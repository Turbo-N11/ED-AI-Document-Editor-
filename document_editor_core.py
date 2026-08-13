import os
import subprocess
import re
import json
import base64
import zipfile
import textwrap
import difflib
import requests
import mimetypes
import shutil
import datetime
import contextlib
import asyncio
from io import BytesIO
from urllib.parse import urlparse, parse_qs, unquote

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None

try:
    from PIL import Image
except ImportError:
    Image = None

# -------------------------------
# Set API Key via environment variable:
#   Windows (cmd): set OPENROUTER_API_KEY=sk-or-v1-xxxxxxxx
#   macOS / Linux: export OPENROUTER_API_KEY=sk-or-v1-xxxxxxxx
# -------------------------------
API_KEY = os.environ.get("OPENROUTER_API_KEY", "")
API_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = "openai/gpt-4o-mini"
VISION_MODEL = "openai/gpt-4o-mini"
# -------------------------------

current_doc = None
current_doc_path = None
original_doc = None
improvement_report = []

# =========================================================
# Enhanced editor infrastructure
# =========================================================

AUTOSAVE_ENABLED = True
AUTOSAVE_DIR = os.path.join(os.getcwd(), "autosave")
IMAGES_DIR = os.path.join(os.getcwd(), "images")
TEMPLATES_DIR = os.path.join(os.getcwd(), "templates")
LEGAL_CHECK_DIR = os.path.join(os.getcwd(), "legal_checks")
IMAGE_INDEX_FILENAME = "image_index.json"

# Autosave is one rolling file for the currently loaded document.
# Example: Read.docx -> autosave/Read_autosave.docx
AUTOSAVE_CURRENT_PATH = None

# Undo / redo history. Each entry is a deep copy of the document state
# immediately before a successful editing command.
UNDO_STACK = []
REDO_STACK = []
MAX_HISTORY = 50


def _clear_history():
    UNDO_STACK.clear()
    REDO_STACK.clear()


def _snapshot_document():
    """Return an independent snapshot of the current document."""
    if current_doc is None:
        return None
    import copy
    return copy.deepcopy(current_doc)


def _push_undo_snapshot(snapshot):
    if snapshot is None:
        return
    UNDO_STACK.append(snapshot)
    if len(UNDO_STACK) > MAX_HISTORY:
        del UNDO_STACK[0]
    REDO_STACK.clear()


def _autosave_after_history_change(reason):
    if current_doc is not None:
        autosave_current_document(reason)


def undo_document():
    global current_doc
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return
    if not UNDO_STACK:
        print("Nothing to undo.")
        return
    import copy
    REDO_STACK.append(copy.deepcopy(current_doc))
    current_doc = UNDO_STACK.pop()
    _autosave_after_history_change("undo")
    print(f"[{document_name()}] Undo applied.")


def redo_document():
    global current_doc
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return
    if not REDO_STACK:
        print("Nothing to redo.")
        return
    import copy
    UNDO_STACK.append(copy.deepcopy(current_doc))
    current_doc = REDO_STACK.pop()
    if len(UNDO_STACK) > MAX_HISTORY:
        del UNDO_STACK[0]
    _autosave_after_history_change("redo")
    print(f"[{document_name()}] Redo applied.")



def ensure_working_folders():
    for folder in (AUTOSAVE_DIR, IMAGES_DIR, TEMPLATES_DIR):
        os.makedirs(folder, exist_ok=True)


def _get_autosave_path():
    """Return the single rolling autosave path for the loaded document."""
    global AUTOSAVE_CURRENT_PATH
    base = os.path.splitext(os.path.basename(current_doc_path or "document.docx"))[0]
    AUTOSAVE_CURRENT_PATH = os.path.join(AUTOSAVE_DIR, f"{base}_autosave.docx")
    return AUTOSAVE_CURRENT_PATH

def _remove_old_autosave_files(keep_path=None):
    """Keep the autosave folder limited to the current rolling autosave."""
    try:
        ensure_working_folders()
        keep_abs = os.path.abspath(keep_path) if keep_path else None
        for name in os.listdir(AUTOSAVE_DIR):
            if not name.lower().endswith(".docx"):
                continue
            path = os.path.abspath(os.path.join(AUTOSAVE_DIR, name))
            if keep_abs and path == keep_abs:
                continue
            try:
                os.remove(path)
            except OSError:
                pass
    except Exception:
        pass

def autosave_current_document(reason="change"):
    """Save every change to one fixed <document>_autosave.docx file."""
    if not AUTOSAVE_ENABLED or current_doc is None:
        return
    try:
        ensure_working_folders()
        out = _get_autosave_path()
        _remove_old_autosave_files(keep_path=out)
        current_doc.save(out)
        print(f"[AUTOSAVE] {os.path.basename(out)}")
    except Exception as e:
        print(f"[AUTOSAVE] Failed: {e}")


def set_run_highlight(run, color="lightBlue"):
    try:
        rPr = run._r.get_or_add_rPr()
        old = rPr.find(qn("w:highlight"))
        if old is not None:
            rPr.remove(old)
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), color)
        rPr.append(hl)
    except Exception:
        pass


def set_cell_shading(cell, fill):
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:fill"), re.sub(r"[^0-9A-Fa-f]", "", fill)[:6])
    except Exception:
        pass


def iter_all_paragraphs(doc):
    for i, p in enumerate(doc.paragraphs):
        yield f"Paragraph {i}", p
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield f"Table {ti}, Row {ri}, Cell {ci}, Paragraph {pi}", p
    for si, sec in enumerate(doc.sections):
        for pi, p in enumerate(sec.header.paragraphs):
            yield f"Header {si}, Paragraph {pi}", p
        for pi, p in enumerate(sec.footer.paragraphs):
            yield f"Footer {si}, Paragraph {pi}", p


def is_heading_paragraph(p):
    kind = identify_paragraph_type(p)
    return kind == "TITLE" or "HEADING" in kind.upper()


def _replace_text_preserve_runs(paragraph, old, new, case_sensitive=False):
    if not paragraph.text:
        return 0
    flags = 0 if case_sensitive else re.IGNORECASE
    pattern = re.compile(re.escape(old), flags)
    changed = 0
    for run in paragraph.runs:
        if pattern.search(run.text or ""):
            run.text = pattern.sub(new, run.text)
            changed += 1
    if changed == 0 and pattern.search(paragraph.text):
        set_paragraph_text(paragraph, pattern.sub(new, paragraph.text))
        changed = 1
    return changed


def _format_runs(runs, spec):
    changed = 0
    for run in runs:
        if spec.get("font_name"):
            run.font.name = spec["font_name"]
            changed += 1
        if spec.get("font_size_pt") is not None:
            run.font.size = Pt(float(spec["font_size_pt"]))
            changed += 1
        if spec.get("font_color"):
            run.font.color.rgb = parse_hex_color(spec["font_color"])
            changed += 1
        if spec.get("bold") is not None:
            run.bold = bool(spec["bold"])
            changed += 1
        if spec.get("italic") is not None:
            run.italic = bool(spec["italic"])
            changed += 1
        if spec.get("underline") is not None:
            run.underline = bool(spec["underline"])
            changed += 1
    return changed


def _apply_paragraph_format(p, spec):
    changed = False
    amap = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    align = str(spec.get("alignment", "")).lower()
    if align in amap:
        p.alignment = amap[align]
        changed = True
    if spec.get("line_spacing") is not None:
        p.paragraph_format.line_spacing = float(spec["line_spacing"])
        changed = True
    if spec.get("space_after_pt") is not None:
        p.paragraph_format.space_after = Pt(float(spec["space_after_pt"]))
        changed = True
    if spec.get("space_before_pt") is not None:
        p.paragraph_format.space_before = Pt(float(spec["space_before_pt"]))
        changed = True
    return changed


def _json_from_ai(raw):
    try:
        return json.loads(raw)
    except Exception:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                pass
    return {}


def _normalize_json_list(value):
    if isinstance(value, list):
        return value
    if isinstance(value, dict):
        return [value]
    return []


def _image_paragraphs(doc):
    return [(location, p) for location, p in iter_all_paragraphs(doc)
            if "drawing" in p._element.xml]


def _image_catalog_from_doc(doc):
    """Return actual embedded-picture objects, not merely paragraphs containing drawings.

    A paragraph can contain a drawing that is not an image (or can contain a
    drawing nested inside runs).  The old implementation treated every such
    paragraph as an image and then tried to read the relationship directly
    from the paragraph XML.  That caused commands such as `replace image img1
    with img3` to report that the image had no replaceable relationship.
    """
    catalog = []
    idx = 0

    for location, p in iter_all_paragraphs(doc):
        # Find actual DrawingML image references.  We deliberately look for
        # r:embed on a:blip elements instead of the generic word "drawing".
        xml = p._element.xml
        rel_ids = re.findall(r'<(?:[A-Za-z0-9_]+:)?blip\b[^>]*?\br:embed="([^"]+)"', xml)

        # Also support the uncommon linked-image form.
        if not rel_ids:
            rel_ids = re.findall(r'<(?:[A-Za-z0-9_]+:)?blip\b[^>]*?\br:link="([^"]+)"', xml)

        for rid in rel_ids:
            try:
                rel = p.part.rels[rid]
                target = str(getattr(rel._target, "partname", ""))
                # Ignore non-image relationships.
                if target and not target.lower().startswith("/word/media/"):
                    continue
            except Exception:
                continue

            idx += 1
            catalog.append({
                "index": idx,
                "location": location,
                "paragraph": p,
                "rel_id": rid,
                "partname": target,
            })

    return catalog


def _find_image_source(name):
    ensure_working_folders()
    name = name.strip().strip('"').strip("'")
    candidates = [name]
    if not os.path.splitext(name)[1]:
        candidates.extend(name + ext for ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp"))
    for candidate in candidates:
        if os.path.isfile(candidate):
            return candidate
        candidate2 = os.path.join(IMAGES_DIR, os.path.basename(candidate))
        if os.path.isfile(candidate2):
            return candidate2
    return None


def _load_image_index():
    ensure_working_folders()
    fn = os.path.join(IMAGES_DIR, IMAGE_INDEX_FILENAME)
    try:
        with open(fn, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def _save_image_index(index):
    ensure_working_folders()
    fn = os.path.join(IMAGES_DIR, IMAGE_INDEX_FILENAME)
    with open(fn, "w", encoding="utf-8") as f:
        json.dump(index, f, ensure_ascii=False, indent=2)


def _extract_embedded_images_to_images_folder(docx_path):
    ensure_working_folders()
    exported = []
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            media = sorted(n for n in z.namelist() if n.startswith("word/media/"))
            for i, filename in enumerate(media, 1):
                ext = os.path.splitext(filename)[1].lower()
                if ext not in (".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif"):
                    continue
                out = os.path.join(IMAGES_DIR, f"img{i}{ext}")
                with open(out, "wb") as f:
                    f.write(z.read(filename))
                exported.append(out)
    except Exception as e:
        print(f"[IMAGES] Extraction failed: {e}")
    return exported


def _index_document_images_with_ai(doc):
    exported = _extract_embedded_images_to_images_folder(current_doc_path) if current_doc_path else []
    catalog = _image_catalog_from_doc(doc)
    old_index = _load_image_index()
    updated = {}
    for pos, item in enumerate(catalog, 1):
        stable = f"img{pos}"
        image_path = exported[pos - 1] if pos <= len(exported) else old_index.get(stable, {}).get("file")
        label = old_index.get(stable, {}).get("label", "Unknown image")
        if image_path and os.path.isfile(image_path):
            try:
                label = chat_with_openrouter(
                    "Identify the main subject of this image in 3-8 words. Return only a concise noun phrase.",
                    system_prompt="You are an image cataloging assistant.",
                    max_tokens=30,
                    stream_to_screen=False,
                    vision_image_path=image_path,
                ).strip().strip('"') or label
            except Exception:
                pass
        updated[stable] = {
            "index": pos,
            "label": label,
            "location": item["location"],
            "partname": item["partname"],
            "file": image_path,
        }
    old_index.update(updated)
    _save_image_index(old_index)
    if updated:
        print("\n=== IMAGE INDEX ===")
        for name, item in updated.items():
            print(f"{name}: {item['label']} [{item['location']}]")
        print(f"Image files: {IMAGES_DIR}")
    return updated


# =========================================================
# Text cleanup
# =========================================================

def clean_chunk(chunk):
    chunk = chunk.replace('**', '')
    chunk = chunk.replace('*', '')
    chunk = re.sub(r'\\\((.*?)\\\)', r'\1', chunk)
    chunk = re.sub(r'\\\[(.*?)\\\]', r'\1', chunk)
    return chunk


def clean_response(text, width=80):
    text = re.sub(r'^(\d+)\.', r'\1)', text, flags=re.MULTILINE)
    text = re.sub(r'^- ', '  * ', text, flags=re.MULTILINE)
    text = text.replace('**', '').replace('*', '')

    wrapped_lines = []
    for line in text.splitlines():
        stripped = line.lstrip()
        indent = len(line) - len(stripped)
        wrapper = textwrap.TextWrapper(width=width, subsequent_indent=' ' * indent)
        wrapped_lines.append(wrapper.fill(stripped))
    text = "\n".join(wrapped_lines)

    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def parse_hex_color(color_hex):
    """Converts a hex string like #FF0000 or color names to RGBColor."""
    color_map = {
        "black": RGBColor(0, 0, 0),
        "white": RGBColor(255, 255, 255),
        "red": RGBColor(255, 0, 0),
        "dark red": RGBColor(139, 0, 0),
        "green": RGBColor(0, 128, 0),
        "dark green": RGBColor(0, 100, 0),
        "blue": RGBColor(0, 0, 255),
        "light blue": RGBColor(173, 216, 230),
        "lightblue": RGBColor(173, 216, 230),
        "navy": RGBColor(0, 0, 128),
        "dark blue": RGBColor(0, 0, 139),
        "gray": RGBColor(128, 128, 128),
        "grey": RGBColor(128, 128, 128),
        "dark gray": RGBColor(169, 169, 169),
        "purple": RGBColor(128, 0, 128),
    }

    if not color_hex:
        return RGBColor(0, 0, 0)

    clean_c = color_hex.strip().lower()
    if clean_c in color_map:
        return color_map[clean_c]

    clean_hex = re.sub(r'[^0-9a-fA-F]', '', color_hex)
    if len(clean_hex) == 6:
        r = int(clean_hex[0:2], 16)
        g = int(clean_hex[2:4], 16)
        b = int(clean_hex[4:6], 16)
        return RGBColor(r, g, b)

    return RGBColor(0, 0, 0)


# =========================================================
# OpenRouter chat
# =========================================================

def chat_with_openrouter(user_input, system_prompt="You are a helpful AI assistant.",
                          max_tokens=250, stream_to_screen=True, 
                          vision_image_path=None, vision_image_bytes=None, vision_mime_type="image/jpeg"):
    if not API_KEY:
        return ("Error: OPENROUTER_API_KEY is not set. Set it in your environment variables.")

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    b64_image = None
    if vision_image_bytes:
        b64_image = base64.b64encode(vision_image_bytes).decode("utf-8")
        mime_type = vision_mime_type
    elif vision_image_path:
        try:
            mime_type, _ = mimetypes.guess_type(vision_image_path)
            if not mime_type or not mime_type.startswith("image/"):
                mime_type = "image/jpeg"
            with open(vision_image_path, "rb") as img_file:
                b64_image = base64.b64encode(img_file.read()).decode("utf-8")
        except Exception as e:
            return f"Error reading image for vision: {e}"

    if b64_image:
        messages_payload = [
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": user_input},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{b64_image}"}
                    }
                ]
            }
        ]
        model_to_use = VISION_MODEL
    else:
        messages_payload = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_input}
        ]
        model_to_use = MODEL

    data = {
        "model": model_to_use,
        "messages": messages_payload,
        "temperature": 0.3,
        "max_tokens": max_tokens,
        "stream": True
    }

    try:
        with requests.post(API_URL, headers=headers, json=data, stream=True) as response:
            if response.status_code != 200:
                return f"Error {response.status_code}: {response.text}"

            buffer = ""
            for line in response.iter_lines():
                if line:
                    decoded = line.decode('utf-8').strip()
                    if decoded.startswith("data: "):
                        decoded = decoded[6:]
                    if decoded == "[DONE]":
                        break
                    try:
                        chunk = json.loads(decoded)
                        delta = chunk.get("choices", [{}])[0].get("delta", {}).get("content", "")
                        delta = clean_chunk(delta)
                        buffer += delta
                        if stream_to_screen:
                            print(delta, end='', flush=True)
                    except json.JSONDecodeError:
                        continue
            if stream_to_screen:
                print()
            return buffer
    except requests.exceptions.RequestException as e:
        return f"Request error: {e}"


# =========================================================
# Document handling (.docx)
# =========================================================

def list_docx_files(folder="."):
    return sorted(
        f for f in os.listdir(folder)
        if f.lower().endswith(".docx") and not f.startswith("~$")
    )


def load_docx(filename):
    global current_doc, current_doc_path, original_doc, AUTOSAVE_CURRENT_PATH

    if Document is None:
        print("python-docx isn't installed. Run: pip install python-docx")
        return

    filename = filename.strip().strip('"')
    if not os.path.isfile(filename):
        candidates = [
            f for f in list_docx_files()
            if f.lower() == filename.lower() or f.lower() == (filename + ".docx").lower()
        ]
        if candidates:
            filename = candidates[0]
        else:
            print(f"Couldn't find '{filename}' in this folder. Try 'list docs'.")
            return

    try:
        current_doc = Document(filename)
        import copy
        original_doc = copy.deepcopy(current_doc)
        current_doc_path = filename
        AUTOSAVE_CURRENT_PATH = None
        _clear_history()
        ensure_working_folders()

        # Create the first rolling autosave immediately after loading.
        if AUTOSAVE_ENABLED:
            autosave_current_document("initial load")
    except Exception as e:
        print(f"Couldn't open '{filename}': {e}")
        return

    print(f"[{filename}] Loaded — {len(current_doc.paragraphs)} paragraphs. Type 'show' to view them.")
    # Image indexing is intentionally NOT performed on load.
    # Run `images` / `index images` only when the user explicitly requests it.


def identify_paragraph_type(paragraph):
    style = ""
    try:
        style = paragraph.style.name or ""
    except Exception:
        pass

    style_low = style.lower()

    if "title" in style_low:
        return "TITLE"
    if "heading" in style_low:
        return style.upper()

    text = paragraph.text.strip()
    if text:
        if len(text) <= 120 and (
            text.isupper()
            or re.match(r'^(article|section|clause|chapter|part)\b', text, re.I)
            or re.match(r'^\d+(?:\.\d+)*[\s.)-]+', text)
        ):
            return "HEADING?"

    return "PARAGRAPH"


def show_paragraphs():
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    print(f"\n=== PARAGRAPHS: [{document_name()}] ===\n")
    shown = False

    for i, p in enumerate(current_doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        kind = identify_paragraph_type(p)
        preview = (text[:100] + "…") if len(text) > 100 else text

        if kind == "PARAGRAPH":
            print(f"[{i}] {preview}")
        else:
            print(f"[{i}] <{kind}> {preview}")

        shown = True

    if not shown:
        print("(document has no non-empty body paragraphs)")

    if current_doc.tables:
        print(f"\n[{document_name()}] Tables: {len(current_doc.tables)}")
        for ti, table in enumerate(current_doc.tables):
            print(f"  [TABLE {ti}] {len(table.rows)} rows × {len(table.columns)} columns")


def set_paragraph_text(paragraph, new_text):
    for run in paragraph.runs[1:]:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def find_and_replace(old, new):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    count = 0
    table_count = 0

    header_footer_count = 0
    for location, p in iter_all_paragraphs(current_doc):
        if old in p.text:
            set_paragraph_text(p, p.text.replace(old, new))
            if location.startswith("Header ") or location.startswith("Footer "):
                header_footer_count += 1
            elif location.startswith("Table "):
                table_count += 1
            else:
                count += 1

    total = count + table_count + header_footer_count
    if total:
        print(f"[{document_name()}] Replaced in {total} location(s).")
        print(f"  Body paragraphs: {count}")
        print(f"  Table paragraphs: {table_count}")
        print(f"  Header/footer paragraphs: {header_footer_count}")
        autosave_current_document("find and replace")
    else:
        print(f'"{old}" was not found in the document.')


def rewrite_paragraph(index, instruction):
    if current_doc is None:
        print("No document loaded. Use: open <filename.docx>")
        return

    paragraphs = current_doc.paragraphs
    if not (0 <= index < len(paragraphs)):
        print(f"No paragraph #{index}. Use 'show' to see valid indices.")
        return

    original = paragraphs[index].text
    if not original.strip():
        print("That paragraph is empty.")
        return

    system_prompt = (
        "You are an expert editor. Rewrite the given paragraph according to the "
        "instruction. Reply with ONLY the rewritten paragraph text — no quotes, "
        "no explanation, no preamble."
    )
    prompt = f"Instruction: {instruction}\n\nParagraph:\n{original}"

    print(f"Rewriting paragraph [{index}]...\n")
    new_text = chat_with_openrouter(prompt, system_prompt=system_prompt, max_tokens=400)
    new_text = new_text.strip().strip('"')

    if new_text.startswith("Error"):
        print(new_text)
        return

    set_paragraph_text(paragraphs[index], new_text)
    print(f"\n\nParagraph [{index}] updated.")
    autosave_current_document("paragraph rewrite")


def save_docx(filename=None):
    if current_doc is None:
        print("No document loaded.")
        return
    if filename is None:
        filename = current_doc_path
    filename = filename.strip().strip('"')
    if not filename.lower().endswith(".docx"):
        filename += ".docx"
    try:
        current_doc.save(filename)
        print(f"[{os.path.basename(filename)}] Saved successfully.")
        globals()["current_doc_path"] = filename
    except Exception as e:
        print(f"Couldn't save: {e}")


# =========================================================
# Feature: Custom Document Editing ("doc edit : ...")
# =========================================================

def doc_edit(instructions):
    """Plan and apply multiple independent formatting/edit requests in one pass."""
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return
    instructions = instructions.strip()
    if not instructions:
        print("Usage: doc edit : change headers to Arial green and highlight every Name light blue")
        return

    context = extract_document_text(current_doc, max_chars=50000)
    system_prompt = """You are a precise Microsoft Word editing planner.
Convert the user's COMPLETE multi-part instruction into JSON only:
{
 "operations": [
  {"type":"format","target":"all|headers|body|tables|table_headers",
   "font_name":null,"font_size_pt":null,"font_color":null,"bold":null,
   "italic":null,"underline":null,"alignment":null,"line_spacing":null,
   "space_after_pt":null,"space_before_pt":null},
  {"type":"highlight","text":"exact text","color":"lightBlue"},
  {"type":"replace","find":"exact existing text","replace":"new text",
   "case_sensitive":false,"scope":"all|body|tables"}
 ]
}
Create one operation for EACH distinct user request. Do not drop parts of a sentence.
Use exact existing text for replacements/highlights when possible. Never invent text.
If an instruction cannot be safely represented, omit only that operation."""
    raw = chat_with_openrouter(
        f"USER INSTRUCTIONS:\n{instructions}\n\nDOCUMENT:\n{context}",
        system_prompt=system_prompt,
        max_tokens=1800,
        stream_to_screen=False,
    ).strip()
    plan = _json_from_ai(raw)
    operations = _normalize_json_list(plan.get("operations"))
    if not operations:
        print("Failed to create a safe multi-operation edit plan.")
        return

    target_count = 0
    for op in operations:
        typ = str(op.get("type", "")).lower()

        if typ == "format":
            target = str(op.get("target", "all")).lower()
            spec = {
                "font_name": op.get("font_name"),
                "font_size_pt": op.get("font_size_pt"),
                "font_color": op.get("font_color"),
                "bold": op.get("bold"),
                "italic": op.get("italic"),
                "underline": op.get("underline"),
                "alignment": op.get("alignment"),
                "line_spacing": op.get("line_spacing"),
                "space_after_pt": op.get("space_after_pt"),
                "space_before_pt": op.get("space_before_pt"),
            }
            if target == "headers":
                targets = [(loc, p) for loc, p in iter_all_paragraphs(current_doc)
                           if is_heading_paragraph(p)]
            elif target == "body":
                targets = [(loc, p) for loc, p in iter_all_paragraphs(current_doc)
                           if not is_heading_paragraph(p)]
            elif target in ("tables", "table_headers"):
                targets = []
                for ti, table in enumerate(current_doc.tables):
                    for ri, row in enumerate(table.rows):
                        if target == "table_headers" and ri != 0:
                            continue
                        for ci, cell in enumerate(row.cells):
                            for p in cell.paragraphs:
                                targets.append((f"Table {ti} R{ri}C{ci}", p))
            else:
                targets = list(iter_all_paragraphs(current_doc))

            for _, p in targets:
                _apply_paragraph_format(p, spec)
                _format_runs(p.runs, spec)
            target_count += len(targets)

        elif typ == "highlight":
            find_text = str(op.get("text", "")).strip()
            if not find_text:
                continue
            color = str(op.get("color") or "lightBlue")
            for _, p in iter_all_paragraphs(current_doc):
                for run in p.runs:
                    if re.search(re.escape(find_text), run.text or "", re.I):
                        set_run_highlight(run, color)
                        target_count += 1

        elif typ == "replace":
            old = str(op.get("find", "")).strip()
            new = str(op.get("replace", ""))
            if not old:
                continue
            scope = str(op.get("scope", "all")).lower()
            if scope == "tables":
                targets = list(iter_table_paragraphs(current_doc))
            elif scope == "body":
                targets = [(f"Paragraph {i}", p) for i, p in enumerate(current_doc.paragraphs)]
            else:
                targets = list(iter_all_paragraphs(current_doc))
            for _, p in targets:
                target_count += _replace_text_preserve_runs(
                    p, old, new, bool(op.get("case_sensitive", False))
                )

    if target_count:
        print(f"[{document_name()}] Applied {len(operations)} requested operation(s) to {target_count} targets.")
        autosave_current_document("doc edit")
    else:
        print("No safe matches were found; no changes were applied.")


# =========================================================
# Feature: Natural-language document editing
# =========================================================

def _natural_prompt_is_likely_edit(text):
    low = text.strip().lower()
    edit_terms = (
        "change ", "replace ", "edit ", "make ", "set ", "format ",
        "colour ", "color ", "highlight ", "bold ", "italic ", "underline ",
        "font ", "heading ", "headings ", "paragraph ", "rewrite ",
        "improve ", "fix ", "remove ", "delete ", "add ", "turn ",
        "convert ", "update ", "rename ", "correct ", "style ",
        "table ", "image ", "picture ", "template ", "spacing ",
        "alignment ", "align ", "size ", "overall formatting"
    )
    question_starts = (
        "who ", "what ", "when ", "where ", "why ", "how ",
        "which ", "whose ", "is ", "are ", "was ", "were ",
        "can you tell me ", "do you know ", "does ", "did ",
        "tell me ", "explain "
    )
    if low.startswith(edit_terms):
        return True
    if any(term in low for term in (" change ", " replace ", " highlight ",
                                    " make all ", " font ", " heading ",
                                    " formatting", " format ", " rewrite ")):
        return True
    if low.startswith(question_starts):
        return False
    return bool(re.search(r'\b(i want|i need|please|make|change|replace|improve)\b', low))


def _natural_edit_plan(instruction):
    context = extract_document_text(current_doc, max_chars=60000)
    system_prompt = r'''You are the execution planner for an AI Microsoft Word document editor.
The user is allowed to describe edits in ordinary language and does NOT know command syntax.
Turn the COMPLETE request into a JSON object with an `operations` array.

Return ONLY valid JSON:
{
  "operations": [
    {"type":"replace","find":"exact existing text","replace":"new text","case_sensitive":false,"scope":"all|body|tables"},
    {"type":"format","target":"all|headers|body|tables|table_headers","font_name":null,"font_size_pt":null,"font_color":null,"bold":null,"italic":null,"underline":null,"alignment":null,"line_spacing":null,"space_after_pt":null,"space_before_pt":null},
    {"type":"word_format","match":"regex","font_color":null,"highlight":null,"font_name":null,"font_size_pt":null,"bold":null,"italic":null,"underline":null,"scope":"all|body|tables"},
    {"type":"table_replace","find":"exact existing text","replace":"new text","table":null,"column_header":null},
    {"type":"rewrite","paragraph":0,"instruction":"..."}
  ],
  "notes":"..."
}

Rules:
1. Preserve every distinct part of the user's request. Never silently drop a request.
2. For exact quoted replacement requests such as change name "Priya" to "Mayank", use `replace`.
3. For requests such as "heading size to 10", use `format` with target `headers` and font_size_pt 10.
4. For "all document text green, then words beginning with T/t blue", use TWO operations in that order: first format target all with green, then word_format with regex `\b[Tt][A-Za-z]*\b` and blue.
5. For "improve overall formatting", analyze the supplied document and create concrete format operations. Prefer headers, body, spacing, alignment and table headers. Do not invent content.
6. Do not use an `improve_format` operation; produce executable operations directly.
7. Use exact existing text for replacements whenever possible. Do not invent an old value.
8. A font size such as 10 means 10 pt.
9. Map common colors to hex: green=008000, blue=0000FF, red=FF0000, black=000000, white=FFFFFF, yellow=FFFF00, orange=FFA500, purple=800080, gray=808080.
10. For natural-language ranges such as "all headings", use the appropriate target instead of enumerating every paragraph.
11. If a request is ambiguous or unsupported, omit only that operation and explain it in `notes`.
'''
    raw = chat_with_openrouter(
        f"USER REQUEST:\n{instruction}\n\nDOCUMENT:\n{context}",
        system_prompt=system_prompt,
        max_tokens=2400,
        stream_to_screen=False,
    ).strip()
    result = _json_from_ai(raw)
    return result if isinstance(result, dict) else {}


def _apply_word_format(paragraph, pattern, spec):
    if not paragraph.runs or not pattern:
        return 0
    try:
        rx = re.compile(pattern)
    except re.error:
        return 0
    changed = 0
    import copy
    for run in list(paragraph.runs):
        text = run.text or ""
        matches = list(rx.finditer(text))
        if not matches:
            continue
        parent = run._r.getparent()
        insert_at = parent.index(run._r)
        pieces = []
        pos = 0
        for m in matches:
            if m.start() > pos:
                pieces.append((text[pos:m.start()], False))
            pieces.append((m.group(0), True))
            pos = m.end()
        if pos < len(text):
            pieces.append((text[pos:], False))
        base_r = run._r
        for piece_text, matched in pieces:
            if not piece_text:
                continue
            new_r = OxmlElement("w:r")
            rpr = base_r.find(qn("w:rPr"))
            if rpr is not None:
                new_r.append(copy.deepcopy(rpr))
            t = OxmlElement("w:t")
            if piece_text.startswith(" ") or piece_text.endswith(" "):
                t.set(qn("xml:space"), "preserve")
            t.text = piece_text
            new_r.append(t)
            parent.insert(insert_at, new_r)
            insert_at += 1
            if matched:
                from docx.text.run import Run
                new_run = Run(new_r, run._parent)
                _format_runs([new_run], spec)
                if spec.get("highlight"):
                    set_run_highlight(new_run, spec["highlight"])
                changed += 1
        parent.remove(base_r)
    return changed


def natural_document_edit(instruction):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return False
    instruction = instruction.strip()
    if not instruction:
        return False
    plan = _natural_edit_plan(instruction)
    operations = _normalize_json_list(plan.get("operations"))
    if not operations:
        print("I couldn't turn that request into safe document edits.")
        return True
    print(f"\n[NATURAL EDIT] {len(operations)} operation(s) planned:")
    for i, op in enumerate(operations, 1):
        typ = str(op.get("type", "")).lower()
        if typ == "replace":
            print(f"  [{i}] Replace {op.get('find')!r} -> {op.get('replace')!r}")
        elif typ == "format":
            target = op.get("target", "all")
            details = []
            for k in ("font_name", "font_size_pt", "font_color", "bold", "italic", "underline", "alignment"):
                if op.get(k) is not None:
                    details.append(f"{k}={op.get(k)}")
            print(f"  [{i}] Format {target}: {', '.join(details) or 'document formatting'}")
        elif typ == "word_format":
            print(f"  [{i}] Format matching words: {op.get('match')!r}")
        elif typ == "table_replace":
            print(f"  [{i}] Table replacement: {op.get('find')!r} -> {op.get('replace')!r}")
        elif typ == "rewrite":
            print(f"  [{i}] Rewrite paragraph {op.get('paragraph')}: {op.get('instruction')}")
        else:
            print(f"  [{i}] {typ}")
    if plan.get("notes"):
        print(f"  Note: {plan.get('notes')}")

    applied = 0
    for op in operations:
        typ = str(op.get("type", "")).lower()
        if typ == "replace":
            old = str(op.get("find", "")).strip()
            new = str(op.get("replace", ""))
            if not old:
                continue
            scope = str(op.get("scope", "all")).lower()
            if scope == "tables":
                targets = list(iter_table_paragraphs(current_doc))
            elif scope == "body":
                targets = [(f"Paragraph {i}", p) for i, p in enumerate(current_doc.paragraphs)]
            else:
                targets = list(iter_all_paragraphs(current_doc))
            for _, p in targets:
                applied += _replace_text_preserve_runs(p, old, new, bool(op.get("case_sensitive", False)))
        elif typ == "format":
            target = str(op.get("target", "all")).lower()
            spec = {k: op.get(k) for k in ("font_name", "font_size_pt", "font_color", "bold", "italic", "underline", "alignment", "line_spacing", "space_after_pt", "space_before_pt")}
            if target == "headers":
                targets = [(loc, p) for loc, p in iter_all_paragraphs(current_doc) if is_heading_paragraph(p)]
            elif target == "body":
                targets = [(loc, p) for loc, p in iter_all_paragraphs(current_doc) if not is_heading_paragraph(p)]
            elif target in ("tables", "table_headers"):
                targets = []
                for ti, table in enumerate(current_doc.tables):
                    for ri, row in enumerate(table.rows):
                        if target == "table_headers" and ri != 0:
                            continue
                        for ci, cell in enumerate(row.cells):
                            for p in cell.paragraphs:
                                targets.append((f"Table {ti} R{ri}C{ci}", p))
            else:
                targets = list(iter_all_paragraphs(current_doc))
            for _, p in targets:
                _apply_paragraph_format(p, spec)
                _format_runs(p.runs, spec)
                applied += 1
        elif typ == "word_format":
            spec = {k: op.get(k) for k in ("font_name", "font_size_pt", "font_color", "bold", "italic", "underline", "highlight")}
            scope = str(op.get("scope", "all")).lower()
            if scope == "tables":
                targets = list(iter_table_paragraphs(current_doc))
            elif scope == "body":
                targets = [(f"Paragraph {i}", p) for i, p in enumerate(current_doc.paragraphs)]
            else:
                targets = list(iter_all_paragraphs(current_doc))
            for _, p in targets:
                applied += _apply_word_format(p, str(op.get("match", "")), spec)
        elif typ == "table_replace":
            old = str(op.get("find", "")).strip()
            new = str(op.get("replace", ""))
            if not old:
                continue
            table_num = op.get("table")
            if table_num is None:
                tables = list(current_doc.tables)
            else:
                try:
                    n = int(table_num)
                    tables = [current_doc.tables[n]] if 0 <= n < len(current_doc.tables) else []
                except (TypeError, ValueError):
                    tables = []
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            applied += _replace_text_preserve_runs(p, old, new, False)
        elif typ == "rewrite":
            try:
                idx = int(op.get("paragraph"))
            except (TypeError, ValueError):
                continue
            if 0 <= idx < len(current_doc.paragraphs):
                original = current_doc.paragraphs[idx].text
                if original.strip():
                    prompt = ("Rewrite this paragraph according to the user's instruction. Return ONLY the rewritten paragraph text.\n\n" + f"Instruction: {op.get('instruction', '')}\n\nParagraph:\n{original}")
                    rewritten = chat_with_openrouter(prompt, system_prompt="You are a professional document editor.", max_tokens=500, stream_to_screen=False).strip()
                    if rewritten and not rewritten.startswith("Error"):
                        set_paragraph_text(current_doc.paragraphs[idx], rewritten)
                        applied += 1
    if applied:
        print(f"[{document_name()}] Natural-language edit applied ({applied} target operation(s)).")
        autosave_current_document("natural-language edit")
    else:
        print("No matching document content was changed.")
    return True


# =========================================================
# Feature 1: Document Inspection
# =========================================================

def inspect_document():
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    print(f"\n=======================================================")
    print(f"       DOCUMENT INSPECTION REPORT: [{document_name()}]")
    print(f"=======================================================\n")

    total_paras = len(current_doc.paragraphs)
    total_tables = len(current_doc.tables)
    total_sections = len(current_doc.sections)
    
    images_found = []
    for i, p in enumerate(current_doc.paragraphs):
        if "drawing" in p._element.xml:
            images_found.append((f"Paragraph {i}", p))
    for ti, table in enumerate(current_doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    if "drawing" in p._element.xml:
                        images_found.append((f"Table {ti}, R{ri}C{ci}, P{pi}", p))

    print("--- [OVERVIEW] ---")
    print(f"Paragraphs : {total_paras}")
    print(f"Tables     : {total_tables}")
    print(f"Sections   : {total_sections}")
    print(f"Images     : {len(images_found)}")
    print()

    print("--- [SECTIONS, HEADERS & FOOTERS] ---")
    for idx, sec in enumerate(current_doc.sections):
        header_text = " ".join([p.text.strip() for p in sec.header.paragraphs if p.text.strip()]) or "(empty)"
        footer_text = " ".join([p.text.strip() for p in sec.footer.paragraphs if p.text.strip()]) or "(empty)"
        print(f"Section {idx}: Page Width={sec.page_width.inches:.2f}in, Height={sec.page_height.inches:.2f}in")
        print(f"  Header: {header_text}")
        print(f"  Footer: {footer_text}")
    print()

    print("--- [HEADINGS & STRUCTURE] ---")
    heading_count = 0
    for i, p in enumerate(current_doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        kind = identify_paragraph_type(p)
        if kind != "PARAGRAPH":
            heading_count += 1
            print(f"  [{i}] <{kind}> {text}")
    if heading_count == 0:
        print("  (No clear structural headings detected)")
    print()

    print("--- [TABLES DETAILS] ---")
    if current_doc.tables:
        for ti, table in enumerate(current_doc.tables):
            print(f"  Table {ti}: {len(table.rows)} Rows x {len(table.columns)} Cols")
    else:
        print("  (No tables present in document)")
    print()

    print("--- [EMBEDDED IMAGES / DRAWINGS] ---")
    if images_found:
        for idx, (loc, p) in enumerate(images_found):
            print(f"  Image [{idx}]: Located at [{loc}]")
    else:
        print("  (No embedded images detected)")
    print("\n=======================================================\n")


# =========================================================
# Feature 2: Table & Image Editing
# =========================================================

def show_tables_detail():
    """Display document tables as clean terminal-friendly grids."""
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return
    if not current_doc.tables:
        print("No tables found in this document.")
        return

    print(f"\n=== TABLES IN [{document_name()}] ===\n")

    # Keep terminal output readable even when a cell contains long text.
    MAX_CELL_WIDTH = 34

    def clean_cell(value):
        value = str(value or "")
        value = re.sub(r"\\s+", " ", value).strip()
        return value.replace("|", "\\|")

    def shorten(value, width=MAX_CELL_WIDTH):
        value = clean_cell(value)
        if len(value) <= width:
            return value
        return value[:max(1, width - 3)] + "..."

    def print_border(widths, left="┌", middle="┬", right="┐",
                     horizontal="─"):
        print(left + middle.join(horizontal * (w + 2) for w in widths) + right)

    def print_row(values, widths):
        cells = [
            f" {shorten(value, width):<{width}} "
            for value, width in zip(values, widths)
        ]
        print("│" + "│".join(cells) + "│")

    for ti, table in enumerate(current_doc.tables):
        rows = [
            [clean_cell(cell.text).replace("\n", " ") for cell in row.cells]
            for row in table.rows
        ]

        if not rows:
            print(f"Table {ti}: (empty)")
            continue

        column_count = max(len(row) for row in rows)
        for row in rows:
            row.extend([""] * (column_count - len(row)))

        # Header + body.
        header = rows[0]
        body = rows[1:]

        widths = []
        for ci in range(column_count):
            longest = max(
                len(clean_cell(row[ci]))
                for row in rows
            )
            widths.append(min(max(longest, 10), MAX_CELL_WIDTH))

        print(
            f"┌─ TABLE {ti} "
            f"({len(rows)} rows × {column_count} columns) "
            f"{'─' * max(0, sum(widths) + 2 * column_count - 9)}┐"
        )
        print_border(widths)

        print_row(header, widths)
        print("├" + "┼".join("─" * (w + 2) for w in widths) + "┤")

        for row in body:
            print_row(row, widths)

        print_border(widths, left="└", middle="┴", right="┘")
        print()


def edit_table_cell(cmd_args):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return
    m = re.match(r'edit\s+table\s+(\d+)\s+row\s+(\d+)\s+col\s+(\d+)\s*:\s*(.+)',
                 cmd_args, re.I)
    if not m:
        print("Usage: edit table <T> row <R> col <C> : <new text>")
        return
    ti, ri, ci, new_text = int(m.group(1)), int(m.group(2)), int(m.group(3)), m.group(4).strip()
    if ti >= len(current_doc.tables) or ri >= len(current_doc.tables[ti].rows) or ci >= len(current_doc.tables[ti].columns):
        print("Table/row/column index out of range.")
        return
    cell = current_doc.tables[ti].rows[ri].cells[ci]
    set_paragraph_text(cell.paragraphs[0], new_text)
    for p in cell.paragraphs[1:]:
        set_paragraph_text(p, "")
    print(f"[{document_name()}] Table {ti}, Row {ri}, Col {ci} updated.")
    autosave_current_document("table cell edit")


def ai_table_edit(instruction):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return
    if not current_doc.tables:
        print("No tables found.")
        return

    tables = []
    for ti, table in enumerate(current_doc.tables):
        tables.append({
            "table": ti,
            "rows": [{"row": ri, "cells": [c.text.strip() for c in row.cells]}
                     for ri, row in enumerate(table.rows)]
        })

    system_prompt = """You are a safe DOCX table-edit planner.
Return JSON only:
{"operations":[
 {"type":"replace_all","table":0,"find":"Low","replace":"High"},
 {"type":"replace_column","table":0,"column_header":"Status","find":"Low","replace":"High"},
 {"type":"set_cell","table":0,"row":1,"col":2,"text":"new text"}
]}
Only use values present in the supplied tables. Do not invent locations.
For requests like 'turn all Low to High', return replace_all.
Preserve table structure unless the user explicitly requests structural changes."""
    raw = chat_with_openrouter(
        f"USER REQUEST:\n{instruction}\n\nTABLE DATA:\n{json.dumps(tables, ensure_ascii=False)}",
        system_prompt=system_prompt,
        max_tokens=1200,
        stream_to_screen=False,
    ).strip()
    plan = _json_from_ai(raw)
    operations = _normalize_json_list(plan.get("operations"))
    applied = 0

    for op in operations:
        typ = str(op.get("type", "")).lower()
        if typ == "replace_all":
            old, new = str(op.get("find", "")), str(op.get("replace", ""))
            if not old:
                continue
            selected = [int(op["table"])] if op.get("table") is not None else range(len(current_doc.tables))
            for ti in selected:
                if 0 <= ti < len(current_doc.tables):
                    for row in current_doc.tables[ti].rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                applied += _replace_text_preserve_runs(p, old, new, False)

        elif typ == "replace_column":
            header = str(op.get("column_header", "")).strip().lower()
            old, new = str(op.get("find", "")), str(op.get("replace", ""))
            selected = [int(op["table"])] if op.get("table") is not None else range(len(current_doc.tables))
            for ti in selected:
                if not (0 <= ti < len(current_doc.tables)):
                    continue
                table = current_doc.tables[ti]
                col = None
                if table.rows:
                    for ci, cell in enumerate(table.rows[0].cells):
                        if cell.text.strip().lower() == header:
                            col = ci
                            break
                if col is not None:
                    for row in table.rows[1:]:
                        for p in row.cells[col].paragraphs:
                            applied += _replace_text_preserve_runs(p, old, new, False)

        elif typ == "set_cell":
            try:
                ti, ri, ci = int(op["table"]), int(op["row"]), int(op["col"])
                cell = current_doc.tables[ti].rows[ri].cells[ci]
                set_paragraph_text(cell.paragraphs[0], str(op.get("text", "")))
                applied += 1
            except Exception:
                pass

    if applied:
        print(f"[{document_name()}] AI table edit changed {applied} target(s).")
        autosave_current_document("AI table edit")
    else:
        print("No safe table matches were found.")


def replace_image_in_doc(cmd_args):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    m = re.match(r'replace\s+image\s+(\d+|img\d+)\s+with\s+(.+)', cmd_args, re.I)
    if not m:
        print("Usage: replace image <N|imgN> with <file or imgN>")
        return
    img_idx = int(re.sub(r"[^0-9]", "", m.group(1)))
    new_img_path = _find_image_source(m.group(2))
    if not new_img_path:
        print(f"Image source '{m.group(2).strip()}' not found in current folder or ./images.")
        return

    images_found = _image_catalog_from_doc(current_doc)
    if not (1 <= img_idx <= len(images_found)):
        print(f"Image {img_idx} not found. Found {len(images_found)} actual embedded images.")
        return

    image_info = images_found[img_idx - 1]
    target = image_info["paragraph"]
    rel_id = image_info["rel_id"]

    try:
        rel = target.part.rels[rel_id]
        with open(new_img_path, "rb") as f:
            new_blob = f.read()

        # Keep the existing relationship/drawing intact.  Replacing the
        # relationship target blob preserves the image's size, crop,
        # alignment, wrapping, and other drawing properties.
        if hasattr(rel._target, "_blob"):
            rel._target._blob = new_blob
        else:
            # python-docx may expose the target as an ImagePart whose blob is
            # read-only through the public API.  Updating the private blob is
            # the least invasive approach and is what the rest of this tool
            # already relies on for exact drawing preservation.
            rel._target._blob = new_blob
    except Exception as e:
        print(f"Could not replace image safely: {e}")
        return

    print(f"[{document_name()}] Replaced img{img_idx} with '{os.path.basename(new_img_path)}'.")
    print("Original width, height, crop, alignment and wrapping were preserved.")
    autosave_current_document("image replacement")
    _index_document_images_with_ai(current_doc)


# =========================================================
# Helper: Extract Embedded Images from DOCX File
# =========================================================

def extract_images_from_docx(docx_path):
    images = []
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            for filename in sorted(z.namelist()):
                if filename.startswith('word/media/'):
                    ext = os.path.splitext(filename)[1].lower()
                    if ext in ('.png', '.jpg', '.jpeg', '.bmp', '.webp'):
                        mime_type = mimetypes.guess_type(filename)[0] or "image/jpeg"
                        images.append((filename, z.read(filename), mime_type))
    except Exception as e:
        print(f"Error reading zip archive from '{docx_path}': {e}")
    return images


# =========================================================
# Feature 3: AI Document Reconstruction (Images or DOCX) + Verification
# =========================================================

def run_reconstruction_verification(parsed_elements):
    print("\n=======================================================")
    print("       RECONSTRUCTION VERIFICATION REPORT (--verify)")
    print("=======================================================\n")

    reconstructed_text = extract_document_text(current_doc)
    element_summary = json.dumps(parsed_elements, ensure_ascii=False)

    verify_prompt = (
        "You are an AI Document Quality Auditor. Compare the reconstructed document text "
        "against the raw vision OCR elements parsed.\n\n"
        f"PARSED OCR ELEMENTS:\n{element_summary[:3000]}\n\n"
        f"FINAL RECONSTRUCTED DOCX TEXT:\n{reconstructed_text[:3000]}\n\n"
        "Provide a concise verification report answering:\n"
        "1. Completeness Score (0-100%)\n"
        "2. Structure Fidelity (Headings, Tables, Paragraphs)\n"
        "3. Text Accuracy & Any missing/doubtful words\n"
        "4. Final Verification Verdict (PASS / PASS WITH WARNINGS / FAIL)"
    )

    chat_with_openrouter(
        verify_prompt,
        system_prompt="You are a strict document verification auditor.",
        max_tokens=800
    )
    print("\n=======================================================\n")


def reconstruct_document(cmd_args):
    verify = False
    args = cmd_args.strip()
    if "--verify" in args.lower():
        verify = True
        args = re.sub(r'--verify', '', args, flags=re.I).strip()

    source_path = args.strip().strip('"')
    if not os.path.isfile(source_path):
        print(f"Source file '{source_path}' not found.")
        return

    ext = os.path.splitext(source_path)[1].lower()
    images_to_process = []
    if ext == ".docx":
        extracted_images = extract_images_from_docx(source_path)
        if not extracted_images:
            print("No embedded images found.")
            return
        for filename, img_bytes, mime in extracted_images:
            images_to_process.append((os.path.basename(filename), img_bytes, mime, None))
    elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".webp"):
        images_to_process.append((os.path.basename(source_path), None, None, source_path))
    else:
        print("Unsupported format. Use a PNG/JPG image or a DOCX containing scanned pages.")
        return

    all_elements = []
    document_title = ""
    first_page_style = {}

    prompt = """Analyze this scanned document page visually and reconstruct its layout.
Return strict JSON:
{
 "title": "...",
 "page": {
   "width_in": 8.27, "height_in": 11.69,
   "margin_top_in": 0.7, "margin_bottom_in": 0.7,
   "margin_left_in": 0.8, "margin_right_in": 0.8
 },
 "elements": [
  {"type":"heading","level":1,"text":"...",
   "style":{"font":"Arial","size_pt":16,"bold":true,"italic":false,
            "underline":false,"color":"#000000","alignment":"center",
            "line_spacing":1.0,"space_before_pt":0,"space_after_pt":8}},
  {"type":"paragraph","text":"...",
   "style":{"font":"Times New Roman","size_pt":11,"bold":false,"italic":false,
            "underline":false,"color":"#000000","alignment":"justify",
            "line_spacing":1.15,"space_after_pt":6}},
  {"type":"table","rows":[["..."]],
   "style":{"font":"Arial","size_pt":10,"header_bold":true,
            "header_fill":"D9EAF7","alignment":"center","border":"single"}}
 ]
}
Rules:
- Extract full text exactly. Never summarize.
- Infer visual formatting from the pixels: approximate font family, point size,
  weight, italics, color, alignment, spacing and table appearance.
- Preserve the visual hierarchy and relative spacing.
- Do not default all text to black/Arial if the image visibly uses something else.
"""

    for idx, (label, img_bytes, mime, filepath) in enumerate(images_to_process, 1):
        print(f"--- Reconstructing page/image {idx}/{len(images_to_process)}: {label} ---")
        raw = chat_with_openrouter(
            prompt,
            system_prompt="You are an expert OCR and visual document reconstruction engine. Return valid JSON only.",
            max_tokens=5000,
            stream_to_screen=False,
            vision_image_path=filepath,
            vision_image_bytes=img_bytes,
            vision_mime_type=mime or "image/jpeg",
        )
        data = _json_from_ai(raw)
        if not data:
            print(f"Failed to parse reconstruction JSON for {label}.")
            continue
        if not document_title:
            document_title = data.get("title", "")
        if not first_page_style:
            first_page_style = data.get("page") or {}
        all_elements.extend(data.get("elements", []))

    if not all_elements:
        print("No structural elements could be reconstructed.")
        return

    new_doc = Document()
    sec = new_doc.sections[0]
    for attr, value in (
        ("page_width", first_page_style.get("width_in")),
        ("page_height", first_page_style.get("height_in")),
        ("top_margin", first_page_style.get("margin_top_in")),
        ("bottom_margin", first_page_style.get("margin_bottom_in")),
        ("left_margin", first_page_style.get("margin_left_in")),
        ("right_margin", first_page_style.get("margin_right_in")),
    ):
        if value is not None:
            try:
                setattr(sec, attr, Inches(float(value)))
            except Exception:
                pass

    if document_title:
        p = new_doc.add_paragraph(style="Title")
        p.add_run(document_title)

    for el in all_elements:
        typ = el.get("type", "paragraph")
        style = el.get("style") or {}

        if typ == "heading":
            level = min(max(int(el.get("level", 1)), 1), 3)
            p = new_doc.add_paragraph(style=f"Heading {level}")
        elif typ == "table":
            rows = el.get("rows", [])
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = new_doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for ri, rowdata in enumerate(rows):
                for ci, value in enumerate(rowdata):
                    cell = table.rows[ri].cells[ci]
                    p = cell.paragraphs[0]
                    _apply_paragraph_format(p, {
                        "alignment": style.get("alignment"),
                        "line_spacing": 1.0,
                        "space_after_pt": 0,
                    })
                    run = p.add_run(str(value))
                    if style.get("font"):
                        run.font.name = style["font"]
                    if style.get("size_pt"):
                        run.font.size = Pt(float(style["size_pt"]))
                    if ri == 0 and style.get("header_bold"):
                        run.bold = True
                    if style.get("header_fill") and ri == 0:
                        set_cell_shading(cell, style["header_fill"])
            continue
        else:
            p = new_doc.add_paragraph()

        run = p.add_run(str(el.get("text", "")))
        if style.get("font"):
            run.font.name = style["font"]
        if style.get("size_pt"):
            run.font.size = Pt(float(style["size_pt"]))
        if style.get("color"):
            run.font.color.rgb = parse_hex_color(style["color"])
        if style.get("bold") is not None:
            run.bold = bool(style["bold"])
        if style.get("italic") is not None:
            run.italic = bool(style["italic"])
        if style.get("underline") is not None:
            run.underline = bool(style["underline"])
        _apply_paragraph_format(p, {
            "alignment": style.get("alignment"),
            "line_spacing": style.get("line_spacing"),
            "space_after_pt": style.get("space_after_pt"),
            "space_before_pt": style.get("space_before_pt"),
        })

    base_name = os.path.splitext(os.path.basename(source_path))[0]
    out_name = f"reconstructed_{base_name}.docx"
    new_doc.save(out_name)
    print(f"\n[OK] Reconstructed '{out_name}' with visual formatting inferred from the source.")

    load_docx(out_name)
    if verify:
        run_reconstruction_verification(all_elements)


# =========================================================
# Extra document commands
# =========================================================

def reset_document():
    global current_doc
    if current_doc is None or original_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    import copy
    current_doc = copy.deepcopy(original_doc)
    print(f"[{document_name()}] Reset to the original loaded document.")
    print("Unsaved edits made since loading have been discarded.")
    autosave_current_document("reset")


def summarize_document():
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    context = extract_document_text(current_doc)
    prompt = (
        f"DOCUMENT: {document_name()}\n\n"
        "Summarize the entire document. Cover the main purpose, parties, "
        "important dates, amounts, obligations, rights, terms, termination, "
        "and other important provisions that are actually present. "
        "Do not invent missing information. Organize the result with clear "
        "headings and cite paragraph/table locations where possible.\n\n"
        + context
    )
    print(f"\n=== SUMMARY: [{document_name()}] ===\n")
    chat_with_openrouter(
        prompt,
        system_prompt="You are a precise document summarization assistant. "
                      "Summarize only the supplied document.",
        max_tokens=1400
    )


def show_all_paragraphs():
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    print(f"\n=== FULL DOCUMENT: [{document_name()}] ===\n")

    for i, p in enumerate(current_doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        kind = identify_paragraph_type(p)
        if kind == "PARAGRAPH":
            print(f"[{i}] {text}")
        else:
            print(f"[{i}] <{kind}> {text}")

    for ti, table in enumerate(current_doc.tables):
        print(f"\n=== TABLE {ti} ({len(table.rows)} rows × {len(table.columns)} columns) ===")
        for ri, row in enumerate(table.rows):
            cells = []
            for ci, cell in enumerate(row.cells):
                cell_text = " ".join(
                    p.text.strip() for p in cell.paragraphs if p.text.strip()
                )
                cells.append(f"[C{ci}] {cell_text}")
            print(f"[R{ri}] " + " | ".join(cells))


def rewrite_paragraph_range(start, end, instruction):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    paragraphs = current_doc.paragraphs
    if start < 0 or end >= len(paragraphs) or start > end:
        print(f"Invalid paragraph range. Valid range: 0-{len(paragraphs)-1}")
        return

    targets = []
    for i in range(start, end + 1):
        if paragraphs[i].text.strip():
            targets.append((i, paragraphs[i].text))

    if not targets:
        print("The selected paragraph range is empty.")
        return

    print(f"[{document_name()}] Rewriting paragraphs {start}-{end}...")
    for i, original in targets:
        system_prompt = (
            "You are an expert document editor. Rewrite the supplied paragraph "
            "according to the instruction. Return ONLY the rewritten paragraph "
            "text. Preserve facts, names, numbers, dates, and meaning."
        )
        prompt = f"Instruction: {instruction}\n\nParagraph [{i}]:\n{original}"
        new_text = chat_with_openrouter(
            prompt,
            system_prompt=system_prompt,
            max_tokens=500,
            stream_to_screen=False
        ).strip().strip('"')

        if new_text.startswith("Error"):
            print(f"Paragraph [{i}] failed: {new_text}")
            continue

        set_paragraph_text(paragraphs[i], new_text)
        print(f"  [OK] Paragraph [{i}] updated.")

    print(f"Finished paragraphs {start}-{end}.")
    autosave_current_document("paragraph range rewrite")


def document_name():
    return os.path.basename(current_doc_path) if current_doc_path else None


def iter_table_paragraphs(doc):
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    yield f"Table {ti}, Row {ri}, Cell {ci}, Paragraph {pi}", p


def extract_document_text(doc, max_chars=60000):
    parts = []
    total = 0

    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt:
            item = f"[Paragraph {i} | {identify_paragraph_type(p)}]\n{txt}\n\n"
            parts.append(item)
            total += len(item)
            if total >= max_chars:
                return "".join(parts)[:max_chars]

    for ti, table in enumerate(doc.tables):
        item = f"[TABLE {ti}]\n"
        for ri, row in enumerate(table.rows):
            row_text = " | ".join(
                cell.text.strip().replace("\n", " ")
                for cell in row.cells
            )
            item += f"[Row {ri}] {row_text}\n"
        item += "\n"
        parts.append(item)
        total += len(item)
        if total >= max_chars:
            break

    return "".join(parts)[:max_chars]


def load_doc_for_compare(filename):
    filename = filename.strip().strip('"')
    if not os.path.isfile(filename):
        candidates = [f for f in list_docx_files()
                      if f.lower() == filename.lower()
                      or f.lower() == (filename + ".docx").lower()]
        if candidates:
            filename = candidates[0]
        else:
            return None, f"Couldn't find '{filename}'."
    try:
        return Document(filename), None
    except Exception as e:
        return None, f"Couldn't open '{filename}': {e}"


def compare_documents(file1, file2):
    doc1, err = load_doc_for_compare(file1)
    if err:
        print(err)
        return
    doc2, err = load_doc_for_compare(file2)
    if err:
        print(err)
        return

    p1 = [f"<{identify_paragraph_type(p)}> {p.text.strip()}" for p in doc1.paragraphs if p.text.strip()]
    p2 = [f"<{identify_paragraph_type(p)}> {p.text.strip()}" for p in doc2.paragraphs if p.text.strip()]

    matcher = difflib.SequenceMatcher(None, p1, p2)
    changes = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag != "equal":
            changes.append({
                "type": tag, "old": p1[i1:i2], "new": p2[j1:j2],
                "old_range": [i1, max(i1, i2 - 1)],
                "new_range": [j1, max(j1, j2 - 1)]
            })

    t1, t2 = [], []
    for table in doc1.tables:
        for row in table.rows:
            t1.append(" | ".join(cell.text.strip() for cell in row.cells))
    for table in doc2.tables:
        for row in table.rows:
            t2.append(" | ".join(cell.text.strip() for cell in row.cells))

    tm = difflib.SequenceMatcher(None, t1, t2)
    table_changes = []
    for tag, i1, i2, j1, j2 in tm.get_opcodes():
        if tag != "equal":
            table_changes.append({
                "type": tag, "old": t1[i1:i2], "new": t2[j1:j2],
                "old_range": [i1, max(i1, i2 - 1)],
                "new_range": [j1, max(j1, j2 - 1)]
            })

    print("\n=== DOCUMENT COMPARISON ===")
    print(f"OLD: [{os.path.basename(file1)}]")
    print(f"NEW: [{os.path.basename(file2)}]")
    print(f"Paragraph changes: {len(changes)} | Table changes: {len(table_changes)}\n")

    if not changes and not table_changes:
        print("No text changes found.")
        return

    for n, change in enumerate(changes, 1):
        print(f"[Change {n}] {change['type'].upper()} "
              f"(old {change['old_range']}, new {change['new_range']})")
        for x in change["old"]:
            print(f"  - {x}")
        for x in change["new"]:
            print(f"  + {x}")
        print()

    for n, change in enumerate(table_changes, 1):
        print(f"[Table Change {n}] {change['type'].upper()}")
        for x in change["old"]:
            print(f"  - {x}")
        for x in change["new"]:
            print(f"  + {x}")
        print()

    payload = json.dumps({"paragraph_changes": changes,
                          "table_changes": table_changes}, ensure_ascii=False)
    print("=== AI EDIT SUMMARY ===\n")
    chat_with_openrouter(
        "Explain only the actual changes in this diff. Group related edits "
        "such as names, dates, amounts, clauses, additions/removals. "
        "Distinguish substantive changes from wording changes.\n\n" + payload,
        system_prompt="You are a precise document comparison assistant. "
                      "Never invent a change not present in the diff.",
        max_tokens=900
    )


def search_document(query):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    query = query.strip().strip('"').strip("'")
    if not query:
        print('Usage: search "notice period"')
        return

    matches = []
    qlow = query.lower()

    def find_sentences(text):
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z0-9"\'])', text.strip())
        found = [s.strip() for s in sentences if qlow in s.lower()]
        return found or ([text.strip()] if qlow in text.lower() else [])

    for i, p in enumerate(current_doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        for sentence in find_sentences(text):
            matches.append((f"Paragraph {i}", sentence))

    for location, p in iter_table_paragraphs(current_doc):
        text = p.text.strip()
        if not text:
            continue
        for sentence in find_sentences(text):
            matches.append((location, sentence))

    if not matches:
        print(f'No matches found for "{query}".')
        return

    print(f'\nFound {len(matches)} match(es) in [{document_name()}]:\n')
    for location, sentence in matches:
        print(f"[{location}]")
        print(f'"{sentence}"\n')


def ask_document(question):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return
    question = question.strip()
    if not question:
        print("Usage: ask: Who is the landlord?")
        return

    context = extract_document_text(current_doc)
    prompt = (
        f"DOCUMENT: {document_name()}\n\nDOCUMENT CONTENT:\n{context}\n\n"
        f"QUESTION: {question}\n\n"
        "Answer only from the supplied document. If the document does not "
        "contain enough information, say so. Cite paragraph/table locations "
        "shown in the supplied content whenever possible."
    )
    print(f"\n[{document_name()}] Answer:\n")
    chat_with_openrouter(
        prompt,
        system_prompt="You are a document Q&A assistant. Ground answers "
                      "strictly in the supplied document. Never invent facts.",
        max_tokens=700
    )


def smart_global_edit(instruction):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    context = extract_document_text(current_doc)
    prompt = (
        "Identify one exact, safe global replacement from this document.\n"
        f"USER INSTRUCTION: {instruction}\n\nDOCUMENT:\n{context}\n\n"
        'Return JSON only: {"find":"exact old text","replace":"new text",'
        '"reason":"short reason"}.\n'
        'Never invent the old value. If ambiguous, return {"error":"ambiguous"}.'
    )
    raw = chat_with_openrouter(
        prompt,
        system_prompt="You identify exact document text for safe global replacement.",
        max_tokens=350,
        stream_to_screen=False
    ).strip()

    try:
        result = json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r'\{.*\}', raw, re.DOTALL)
        try:
            result = json.loads(m.group(0)) if m else {}
        except Exception:
            result = {}

    old = str(result.get("find", "")).strip()
    new = str(result.get("replace", ""))
    if result.get("error") or not old:
        print("I couldn't identify one unambiguous replacement.")
        print('Example: replace the owner name "Rahul Sharma" everywhere with "John Smith"')
        return

    print(f"\n[{document_name()}]")
    print(f'Proposed: "{old}" -> "{new}"')
    print(f"Reason: {result.get('reason', 'global replacement')}")
    if input("Apply this change? [y/N]: ").strip().lower() != "y":
        print("Change cancelled.")
        return
    find_and_replace(old, new)


def improvements():
    global improvement_report

    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    context = extract_document_text(current_doc)
    prompt = (
        f"DOCUMENT NAME: {document_name()}\n\n"
        "First determine what this document is about from its actual content. "
        "Then review it relative to that purpose. Do not invent facts.\n\n"
        "Return JSON ONLY with this structure:\n"
        '{"document_purpose":"...",'
        '"improvements":[{"title":"...","category":"formatting|clarity|consistency|'
        'missing_term|structure|risk|other","location":"Paragraph 3 or Table 0",'
        '"problem":"...","recommendation":"...",'
        '"replacement_text":"...","action":"edit|add|format|review"}]}\n\n'
        "DOCUMENT:\n" + context
    )

    raw = chat_with_openrouter(
        prompt,
        system_prompt="You are a careful document analyst. Ground every finding "
                      "in the supplied document and return valid JSON only.",
        max_tokens=1800,
        stream_to_screen=False
    ).strip()

    try:
        result = json.loads(raw)
    except Exception:
        m = re.search(r'\{.*\}', raw, re.DOTALL)
        try:
            result = json.loads(m.group(0)) if m else {}
        except Exception:
            result = {}

    improvement_report = result.get("improvements", [])
    purpose = result.get("document_purpose", "Unable to determine from the document.")

    print(f"\n=== IMPROVEMENTS: [{document_name()}] ===")
    print(f"Document purpose: {purpose}\n")

    if not improvement_report:
        print("No specific improvements were identified.")
        return

    for i, item in enumerate(improvement_report, 1):
        print(f"[{i}] {item.get('title', 'Improvement')}")
        print(f"    Category: {item.get('category', 'other')}")
        print(f"    Location: {item.get('location', 'Not specified')}")
        print(f"    Problem: {item.get('problem', '')}")
        print(f"    Recommendation: {item.get('recommendation', '')}")
        if item.get("replacement_text"):
            print(f"    Proposed text: {item.get('replacement_text')}")
        print()


def _parse_improvement_selection(spec, total):
    if not spec:
        return list(range(1, total + 1))

    selected = set()
    for part in spec.split(","):
        part = part.strip()
        if re.fullmatch(r"\d+", part):
            n = int(part)
            if 1 <= n <= total:
                selected.add(n)
        elif re.fullmatch(r"\d+\s*[-–]\s*\d+", part):
            a, b = re.split(r"\s*[-–]\s*", part)
            a, b = int(a), int(b)
            if a > b:
                a, b = b, a
            selected.update(n for n in range(a, b + 1) if n <= total)

    return sorted(selected)


def apply_improvements(selection=None):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    if not improvement_report:
        print("No improvement report is loaded.")
        print("Run 'improvements' first, then use 'improvements --apply'.")
        return

    selected = _parse_improvement_selection(selection, len(improvement_report))
    if not selected:
        print("No valid improvement numbers selected.")
        return

    print(f"\n[{document_name()}] Applying improvements: {', '.join(map(str, selected))}")
    applied = 0

    for number in selected:
        item = improvement_report[number - 1]
        action = str(item.get("action", "")).lower()
        replacement = str(item.get("replacement_text", "")).strip()
        location = str(item.get("location", ""))

        if action == "format":
            print(f"  [{number}] Formatting recommendation requires 'improve format'.")
            continue

        if action == "edit" and replacement:
            m = re.search(r'paragraph\s+(\d+)', location, re.I)
            if m:
                idx = int(m.group(1))
                if 0 <= idx < len(current_doc.paragraphs):
                    set_paragraph_text(current_doc.paragraphs[idx], replacement)
                    applied += 1
                    print(f"  [{number}] Applied to paragraph {idx}.")
                    continue

            m = re.search(
                r'table\s+(\d+).*?row\s+(\d+).*?(?:cell|column)\s+(\d+)',
                location, re.I
            )
            if m:
                ti, ri, ci = map(int, m.groups())
                try:
                    cell = current_doc.tables[ti].rows[ri].cells[ci]
                    if cell.paragraphs:
                        set_paragraph_text(cell.paragraphs[0], replacement)
                        applied += 1
                        print(f"  [{number}] Applied to table {ti}, row {ri}, cell {ci}.")
                        continue
                except Exception:
                    pass

        print(f"  [{number}] Could not safely auto-apply: {item.get('title', '')}")

    print(f"\nApplied {applied} improvement(s). Type 'save' to save them.")


def improve_format():
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return

    print(f"\n=== FORMAT REVIEW: [{document_name()}] ===\n")
    context = extract_document_text(current_doc)

    prompt = (
        "Analyze the document's formatting and return JSON ONLY:\n"
        '{"changes":[{"type":"heading|body|spacing|alignment|table|'
        'numbering|page_layout|other","location":"...","action":"..."}]}\n\n'
        "Create a practical formatting plan. Preserve all substantive text "
        "and facts.\n\n" + context
    )

    raw = chat_with_openrouter(
        prompt,
        system_prompt="You are a DOCX formatting specialist. Return valid JSON only.",
        max_tokens=1000,
        stream_to_screen=False
    ).strip()

    try:
        result = json.loads(raw)
    except Exception:
        m = re.search(r'\{.*\}', raw, re.DOTALL)
        try:
            result = json.loads(m.group(0)) if m else {}
        except Exception:
            result = {}

    changes = result.get("changes", [])
    if not changes:
        print("No formatting changes were identified.")
        return

    print("Proposed formatting changes:\n")
    for i, item in enumerate(changes, 1):
        print(f"[{i}] {item.get('type', 'other').upper()} — "
              f"{item.get('location', 'document')}")
        print(f"    {item.get('action', '')}")

    if input("\nApply these formatting changes? [y/N]: ").strip().lower() != "y":
        print("Formatting changes cancelled.")
        return

    _apply_formatting_changes(changes)


def _get_style_safe(doc, style_name):
    """Return a Word style when available; never fail on custom/missing styles."""
    try:
        return doc.styles[style_name]
    except (KeyError, ValueError, AttributeError):
        return None


def _paragraph_number_from_location(location):
    m = re.search(r'paragraph\s+(\d+)', str(location), re.IGNORECASE)
    return int(m.group(1)) if m else None


def _table_number_from_location(location):
    m = re.search(r'table\s+(\d+)', str(location), re.IGNORECASE)
    return int(m.group(1)) if m else None


def _set_paragraph_style_safe(p, style_name):
    style = _get_style_safe(current_doc, style_name)
    if style is not None:
        try:
            p.style = style
            return True
        except Exception:
            pass
    return False


def _apply_formatting_changes(changes):
    """Apply the actual AI-reviewed changes without assuming built-in styles exist.

    Some DOCX files (especially files created by other editors) do not contain a
    style literally named 'Normal'.  Direct formatting is therefore used as the
    safe fallback instead of aborting the entire operation.
    """
    # Safe defaults. Do not crash if the document has no built-in Normal style.
    body_style = _get_style_safe(current_doc, "Normal")
    if body_style is not None:
        try:
            body_style.font.name = "Times New Roman"
            body_style.font.size = Pt(11)
            body_style.font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = _get_style_safe(current_doc, style_name)
        if style is not None:
            try:
                style.font.name = "Times New Roman"
                style.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass

    # Apply each proposed change at its requested location.
    applied = 0
    for item in changes:
        typ = str(item.get("type", "other")).lower()
        location = str(item.get("location", ""))
        action = str(item.get("action", "")).lower()

        pnum = _paragraph_number_from_location(location)
        if pnum is not None and 0 <= pnum < len(current_doc.paragraphs):
            p = current_doc.paragraphs[pnum]
            if typ == "heading":
                # The review normally says Heading 1/2 in the action.
                if "heading 3" in action:
                    _set_paragraph_style_safe(p, "Heading 3")
                elif "heading 2" in action:
                    _set_paragraph_style_safe(p, "Heading 2")
                else:
                    _set_paragraph_style_safe(p, "Heading 1")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                applied += 1
                continue

            if typ == "body":
                if not _set_paragraph_style_safe(p, "Normal"):
                    # Fallback for documents without a Normal style.
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                applied += 1
                continue

            if typ == "alignment":
                if "center" in action:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif "right" in action:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif "justify" in action:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                applied += 1
                continue

            if typ == "spacing":
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                applied += 1
                continue

        if typ == "table":
            tnum = _table_number_from_location(location)
            if tnum is not None and 0 <= tnum < len(current_doc.tables):
                table = current_doc.tables[tnum]
                # Treat the first row as a header when the review asks for headers.
                if "header" in action or "header" in typ:
                    if table.rows:
                        for cell in table.rows[0].cells:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.bold = True
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            p.paragraph_format.space_after = Pt(3)
                            p.paragraph_format.line_spacing = 1.0
                applied += 1
                continue

    # Normalize only paragraphs that were explicitly reviewed as body/heading.
    # Do not recolor/reformat the entire document indiscriminately.
    print(f"[{document_name()}] Formatting improvements applied: {applied}/{len(changes)} change(s).")
    if applied < len(changes):
        print("Some changes could not be mapped safely to the document.")
    autosave_current_document()



def list_templates():
    ensure_working_folders()
    files = sorted(f for f in os.listdir(TEMPLATES_DIR) if f.lower().endswith(".docx"))
    print("\n".join(files) if files else "No templates found in ./templates.")


def format_from_template(cmd_args):
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return
    m = re.match(r'^format\s+--([^\s]+)(?:\s+(.+))?$', cmd_args.strip(), re.I)
    if not m:
        print("Usage: format --Template_name [optional instructions]")
        return

    template_name = m.group(1)
    extra = (m.group(2) or "").strip()
    ensure_working_folders()
    template_path = None
    for f in os.listdir(TEMPLATES_DIR):
        if os.path.splitext(f)[0].lower() == template_name.lower() or f.lower() == template_name.lower():
            template_path = os.path.join(TEMPLATES_DIR, f)
            break
    if not template_path:
        print(f"Template '{template_name}' not found in ./templates.")
        return

    try:
        template = Document(template_path)

        # Transfer the main Word styles without replacing the current document content.
        for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5"):
            try:
                s = template.styles[style_name]
                d = current_doc.styles[style_name]
                d.font.name = s.font.name
                d.font.size = s.font.size
                d.font.bold = s.font.bold
                d.font.italic = s.font.italic
                d.font.underline = s.font.underline
                try:
                    d.font.color.rgb = s.font.color.rgb
                except Exception:
                    pass
            except Exception:
                pass

        # Assign semantic roles to current content.
        for p in current_doc.paragraphs:
            kind = identify_paragraph_type(p)
            if kind == "TITLE":
                p.style = "Title"
            elif "HEADING" in kind:
                mlevel = re.search(r"(\d+)", kind)
                p.style = f"Heading {min(int(mlevel.group(1)) if mlevel else 1, 5)}"
            else:
                p.style = "Normal"

        for table in current_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.style = "Normal"

        # Copy page geometry/margins from the template.
        if template.sections:
            ts = template.sections[0]
            for sec in current_doc.sections:
                sec.page_width = ts.page_width
                sec.page_height = ts.page_height
                sec.top_margin = ts.top_margin
                sec.bottom_margin = ts.bottom_margin
                sec.left_margin = ts.left_margin
                sec.right_margin = ts.right_margin
                sec.header_distance = ts.header_distance
                sec.footer_distance = ts.footer_distance

        print(f"[{document_name()}] Applied template '{os.path.basename(template_path)}'.")
        if extra:
            doc_edit(extra)
        else:
            autosave_current_document("template formatting")
    except Exception as e:
        print(f"Could not apply template: {e}")


def index_images_command():
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return
    _index_document_images_with_ai(current_doc)


# =========================================================
# Legal document verification (read-only review)
# =========================================================

LEGAL_OFFICIAL_SUFFIXES = (
    ".gov", ".gov.", ".gov.in", ".nic.in", ".gov.uk", ".gov.au",
    ".govt.nz", ".gc.ca", ".gouv.fr", ".gov.sg", ".gov.za",
    ".govt.nz", ".europa.eu", ".legislation.gov.uk"
)


def _legal_host_is_official(url, suggested_domains=None):
    try:
        host = (urlparse(url).hostname or "").lower().rstrip(".")
    except Exception:
        return False
    if not host or not host.startswith(("http://", "https://")) and "://" not in url:
        # host has already been parsed; this branch is intentionally permissive below.
        pass
    if any(host.endswith(suffix) for suffix in LEGAL_OFFICIAL_SUFFIXES):
        return True
    for domain in (suggested_domains or []):
        d = str(domain).lower().strip().replace("https://", "").replace("http://", "").rstrip("/")
        if d and (host == d or host.endswith("." + d)):
            return True
    return False


def _extract_search_url(href):
    """Extract a real URL from common Bing/DDG result links."""
    if not href:
        return None
    href = unquote(href)
    if href.startswith("//"):
        href = "https:" + href
    if href.startswith("/"):
        try:
            q = parse_qs(urlparse(href).query)
            candidate = q.get("q", [None])[0] or q.get("url", [None])[0]
            if candidate:
                return candidate
        except Exception:
            pass
    if href.startswith("http://") or href.startswith("https://"):
        return href
    return None


def _known_official_legal_urls(document_type, jurisdiction):
    """Return high-confidence official starting points when search engines fail."""
    dtype = str(document_type or "").lower()
    j = str(jurisdiction or "").lower()
    urls = []

    # Maharashtra residential tenancy / leave-and-license research.
    if "maharashtra" in j and any(k in dtype for k in (
        "lease", "rental", "rent", "tenancy", "licen", "property"
    )):
        urls.extend([
            "https://www.indiacode.nic.in/handle/123456789/15817?locale=en",
            "https://rentcontrolact.maharashtra.gov.in/Index.aspx",
            "https://lj.maharashtra.gov.in/en/document-category/act-list/page/14/",
            "https://www.indiacode.nic.in/handle/123456789/2190?view_type=browse",
        ])

    # Generic India fallback for central legislation.
    if "india" in j or "indian" in j:
        urls.extend([
            "https://www.indiacode.nic.in/handle/123456789/2190?view_type=browse",
        ])

    # Preserve order and remove duplicates.
    result = []
    for u in urls:
        if u not in result:
            result.append(u)
    return result


def _official_web_search(query, suggested_domains=None, limit=6):
    """Search the public web and return only likely official legal sources.

    Search-engine HTML changes frequently, so this function uses multiple
    extraction patterns and is supplemented by known official starting points.
    """
    headers = {"User-Agent": "Mozilla/5.0 AI-Document-Editor/20"}
    urls = []

    def add_candidates(hrefs):
        for href in hrefs:
            u = _extract_search_url(href)
            if u and _legal_host_is_official(u, suggested_domains) and u not in urls:
                urls.append(u)

    # Bing.
    try:
        r = requests.get(
            "https://www.bing.com/search",
            params={"q": query, "count": 10},
            headers=headers,
            timeout=12,
        )
        if r.ok:
            add_candidates(re.findall(r'<a[^>]+href=["\'](https?://[^"\']+)["\']', r.text, re.I))
            add_candidates(re.findall(r'<li class="b_algo".*?<a[^>]+href=["\']([^"\']+)["\']', r.text, re.I | re.S))
    except Exception:
        pass

    # DuckDuckGo HTML.
    if len(urls) < limit:
        try:
            r = requests.get(
                "https://html.duckduckgo.com/html/",
                params={"q": query},
                headers=headers,
                timeout=12,
            )
            if r.ok:
                add_candidates(re.findall(r'class="result__a"[^>]+href=["\']([^"\']+)["\']', r.text, re.I | re.S))
                add_candidates(re.findall(r'href=["\'](https?://[^"\']+)["\']', r.text, re.I))
        except Exception:
            pass

    return urls[:limit]


def _strip_web_text(html):
    """Convert fetched HTML into compact text suitable for AI review."""
    html = re.sub(r"<script\b[^>]*>.*?</script>", " ", html, flags=re.I | re.S)
    html = re.sub(r"<style\b[^>]*>.*?</style>", " ", html, flags=re.I | re.S)
    html = re.sub(r"<noscript\b[^>]*>.*?</noscript>", " ", html, flags=re.I | re.S)
    html = re.sub(r"<[^>]+>", " ", html)
    html = unquote(html)
    html = re.sub(r"&nbsp;", " ", html, flags=re.I)
    html = re.sub(r"&amp;", "&", html, flags=re.I)
    html = re.sub(r"&quot;", '"', html, flags=re.I)
    html = re.sub(r"&#39;", "'", html, flags=re.I)
    html = re.sub(r"\s+", " ", html)
    return html.strip()


def _fetch_official_legal_source(url, max_chars=12000):
    headers = {"User-Agent": "Mozilla/5.0 AI-Document-Editor/20"}
    try:
        r = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
        if not r.ok:
            return None
        final_url = r.url
        if not _legal_host_is_official(final_url):
            return None
        text = _strip_web_text(r.text)
        if len(text) < 200:
            return None
        return {
            "url": final_url,
            "title": re.sub(r"\s+", " ", (re.search(r"<title[^>]*>(.*?)</title>", r.text, re.I | re.S) or ["", ""])[1]).strip(),
            "text": text[:max_chars],
        }
    except Exception:
        return None


def _legal_classify_document(context):
    prompt = (
        "Classify the supplied document for a legal verification workflow. "
        "Do not give legal advice and do not invent laws. Return JSON only.\n\n"
        "Return exactly: {\"document_type\":\"...\",\"subtype\":\"...\","
        "\"jurisdiction\":\"...\",\"jurisdiction_confidence\":0.0,"
        "\"official_domains\":[\"...\"],\"search_queries\":[\"...\"]}.\n"
        "document_type should be a practical category such as lease, employment "
        "agreement, NDA, service agreement, sale agreement, affidavit, legal notice, "
        "power of attorney, will, petition, privacy policy, loan agreement, or other. "
        "Jurisdiction must be stated only when supported by the document. "
        "Search queries should target official legislation, regulations, courts, or "
        "government guidance relevant to this exact document type and jurisdiction. "
        "Prefer 3-5 concise queries.\n\nDOCUMENT:\n" + context[:50000]
    )
    raw = chat_with_openrouter(
        prompt,
        system_prompt="You classify legal documents and plan authoritative-source research. Output valid JSON only.",
        max_tokens=900,
        stream_to_screen=False,
    ).strip()
    try:
        return json.loads(raw)
    except Exception:
        match = re.search(r"\{.*\}", raw, re.S)
        if match:
            try:
                return json.loads(match.group(0))
            except Exception:
                pass
    return {}


def _legal_build_review(document_context, classification, sources, quick=False):
    """Build a conservative legal review, then run a separate evidence audit.

    The first pass identifies candidate issues. The second pass is deliberately
    stricter: a statutory/legal finding survives only when the supplied official
    source actually supports the proposition. Internal drafting inconsistencies
    are kept separate and do not require a statute.
    """
    source_text = []
    for i, src in enumerate(sources, 1):
        source_text.append(
            f"SOURCE {i}\nTITLE: {src.get('title','')}\nURL: {src['url']}\n"
            f"CONTENT: {src['text']}"
        )
    source_blob = "\n\n==============================\n\n".join(source_text)
    level = "Do a concise first-pass check" if quick else "Perform a thorough cross-check"

    candidate_prompt = (
        "You are an AI legal-document review assistant. This is a source-backed "
        "review, not a determination of legal validity. " + level + ".\n\n"
        "Document classification:\n" + json.dumps(classification, ensure_ascii=False) + "\n\n"
        "DOCUMENT:\n" + document_context[:55000] + "\n\n"
        "AUTHORITATIVE WEB SOURCES:\n" + source_blob[:65000] + "\n\n"
        "CRITICAL EVIDENCE RULES:\n"
        "1. Use ONLY the supplied source content for statutory/regulatory/legal-rule claims.\n"
        "2. A URL or source title by itself is NOT evidence.\n"
        "3. Never invent a statute, section, rule, deadline, penalty, requirement, or case.\n"
        "4. Before calling something a CONFIRMED CONFLICT, identify the exact source "
        "provision and explain the direct conflict with the document.\n"
        "5. If the supplied source text does not contain enough evidence, classify it as "
        "POTENTIAL ISSUE or UNVERIFIED, not CONFIRMED CONFLICT.\n"
        "6. Do not infer a legal requirement merely because it is common practice.\n"
        "7. A missing clause is not automatically illegal.\n"
        "8. Internal contradictions are DRAFTING/CONSISTENCY ISSUES unless a source-backed "
        "legal rule is also established.\n"
        "9. Do not rely on general model knowledge when the supplied sources are silent.\n"
        "10. Every source-backed finding MUST contain source_number, provision_or_section, "
        "and evidence_excerpt copied from the supplied source (maximum 30 words).\n"
        "11. If no source-backed finding can be proven, say so explicitly.\n\n"
        "Return JSON ONLY with this shape:\n"
        '{"overall_status":"...","findings":[{"category":"CONFIRMED CONFLICT|POTENTIAL ISSUE|MISSING/UNCLEAR|DRAFTING/CONSISTENCY ISSUE",'
        '"severity":"HIGH|MEDIUM|LOW","title":"...","location":"...","explanation":"...",'
        '"source_number":1,"provision_or_section":"...","evidence_excerpt":"..."}],'
        '"possible_gaps":["..."],"disclaimer":"..."}'
    )

    raw = chat_with_openrouter(
        candidate_prompt,
        system_prompt=(
            "You perform conservative, source-grounded legal document review. "
            "Never invent legal rules. Output valid JSON only."
        ),
        max_tokens=3200 if not quick else 2000,
        stream_to_screen=False,
    ).strip()

    try:
        candidate = json.loads(raw)
    except Exception:
        m = re.search(r"\{.*\}", raw, re.S)
        try:
            candidate = json.loads(m.group(0)) if m else {}
        except Exception:
            candidate = {}

    if not candidate:
        return "No structured legal review could be produced."

    findings = candidate.get("findings") or []
    # Evidence audit: independently verify each source-backed finding.
    legal_findings = [
        f for f in findings
        if f.get("category") in ("CONFIRMED CONFLICT", "POTENTIAL ISSUE", "MISSING/UNCLEAR")
    ]

    verified = []
    if legal_findings:
        audit_items = []
        for idx, f in enumerate(legal_findings, 1):
            sn = f.get("source_number")
            src = sources[int(sn)-1] if str(sn).isdigit() and 1 <= int(sn) <= len(sources) else None
            audit_items.append(
                f"FINDING {idx}\n"
                f"CLAIM: {f.get('explanation','')}\n"
                f"PROVISION: {f.get('provision_or_section','')}\n"
                f"EXCERPT CLAIMED: {f.get('evidence_excerpt','')}\n"
                f"SOURCE TITLE: {src.get('title','') if src else 'INVALID SOURCE'}\n"
                f"SOURCE URL: {src.get('url','') if src else ''}\n"
                f"SOURCE CONTENT: {src.get('text','')[:14000] if src else 'NO SOURCE'}"
            )

        audit_prompt = (
            "Audit the following proposed legal findings against the supplied official source text.\n"
            "Return JSON ONLY: {\"verified\":[{\"finding_number\":1,\"status\":\"SUPPORTED|NOT_SUPPORTED|PARTIALLY_SUPPORTED\","
            "\"reason\":\"...\",\"corrected_provision\":\"...\"}],\"general_note\":\"...\"}\n\n"
            "Rules:\n"
            "- SUPPORTED means the source text actually supports the legal proposition, not merely that the source is relevant.\n"
            "- NOT_SUPPORTED means the citation does not establish the claim; it must be removed as a confirmed legal finding.\n"
            "- PARTIALLY_SUPPORTED means the source supports only part of the claim; narrow it.\n"
            "- Never use outside knowledge.\n\n" + "\n\n====================\n\n".join(audit_items)
        )
        audit_raw = chat_with_openrouter(
            audit_prompt,
            system_prompt="You are an evidence auditor. Verify only what is explicitly supported by the supplied source text.",
            max_tokens=2200 if not quick else 1300,
            stream_to_screen=False,
        ).strip()
        try:
            audit = json.loads(audit_raw)
        except Exception:
            am = re.search(r"\{.*\}", audit_raw, re.S)
            try:
                audit = json.loads(am.group(0)) if am else {}
            except Exception:
                audit = {}

        audit_by_num = {int(x.get("finding_number")): x for x in audit.get("verified", [])
                        if str(x.get("finding_number", "")).isdigit()}
        for i, f in enumerate(legal_findings, 1):
            a = audit_by_num.get(i, {})
            status = str(a.get("status", "NOT_SUPPORTED")).upper()
            if status == "SUPPORTED":
                f["evidence_status"] = "SUPPORTED"
                verified.append(f)
            elif status == "PARTIALLY_SUPPORTED":
                f["evidence_status"] = "PARTIALLY_SUPPORTED"
                f["severity"] = "LOW" if f.get("severity") == "HIGH" else f.get("severity", "LOW")
                f["explanation"] = a.get("reason") or f.get("explanation", "")
                if a.get("corrected_provision"):
                    f["provision_or_section"] = a["corrected_provision"]
                verified.append(f)
            # NOT_SUPPORTED findings are intentionally dropped.

    # Keep drafting findings without pretending they are statutory conflicts.
    drafting = [f for f in findings if f.get("category") == "DRAFTING/CONSISTENCY ISSUE"]
    final_findings = verified + drafting
    candidate["findings"] = final_findings
    candidate["overall_status"] = (
        "Issues require review" if final_findings else "No verified source-backed conflict found"
    )

    # Render a stable human-readable report from the verified JSON.
    lines = [
        "Document Review Report",
        "",
        f"Document Type: {classification.get('document_type','Unknown')}",
        f"Jurisdiction: {classification.get('jurisdiction','Unknown')}",
        f"Overall Status: {candidate.get('overall_status','Review')}",
        "",
        "---",
        "",
        "### Findings:",
    ]
    categories = [
        "CONFIRMED CONFLICT", "POTENTIAL ISSUE", "MISSING/UNCLEAR", "DRAFTING/CONSISTENCY ISSUE"
    ]
    for cat in categories:
        group = [f for f in final_findings if f.get("category") == cat]
        if not group:
            continue
        lines += [f"\n#### {cat}:", ""]
        for n, f in enumerate(group, 1):
            lines += [
                f"{n}. {f.get('title','Untitled')} [{f.get('severity','MEDIUM')}]:",
                f"   - Location: {f.get('location','Not specified')}",
                f"   - {f.get('explanation','')}",
            ]
            if f.get("source_number") and f.get("evidence_status") in ("SUPPORTED", "PARTIALLY_SUPPORTED"):
                lines += [
                    f"   - Source: SOURCE {f.get('source_number')}",
                    f"   - Provision: {f.get('provision_or_section','Not specified')}",
                    f"   - Evidence: {f.get('evidence_excerpt','Not available')}",
                ]

    gaps = candidate.get("possible_gaps") or []
    if gaps:
        lines += ["\n### Possible Gaps:", ""]
        lines += [f"- {g}" for g in gaps]

    lines += [
        "\n### Evidence Policy:",
        "- Source-backed findings are shown only after an independent evidence audit.",
        "- Unsupported statutory citations are excluded rather than presented as facts.",
        "- Internal drafting inconsistencies are reported separately from legal conflicts.",
        "",
        "Disclaimer: This is an AI-assisted document review, not legal advice or a determination of legal validity."
    ]
    return "\n".join(lines)


def check_legal_document(cmd_args=""):
    """Read-only legal verification against fetched official sources."""
    if current_doc is None:
        print("No document loaded. Use: load <filename.docx>")
        return
    if not API_KEY:
        print("OPENROUTER_API_KEY is not set. Legal checking requires the AI model.")
        return

    args = cmd_args.strip()
    quick = bool(re.search(r"(?:^|\s)--quick(?:\s|$)", args, re.I))
    jm = re.search(r"--jurisdiction\s+(.+?)(?=\s+--|$)", args, re.I)
    forced_jurisdiction = jm.group(1).strip() if jm else None

    os.makedirs(LEGAL_CHECK_DIR, exist_ok=True)
    print("\n╭─────────────────────────────────────────────────────────────╮")
    print("│                    LEGAL DOCUMENT CHECK                   │")
    print("╰─────────────────────────────────────────────────────────────╯\n")
    print(f"Document: {document_name()}")
    print("Reading document and identifying legal document type...\n")

    context = extract_document_text(current_doc, max_chars=60000)
    classification = _legal_classify_document(context)
    if not classification:
        print("Could not reliably classify the document. No legal web check was performed.")
        return

    if forced_jurisdiction:
        classification["jurisdiction"] = forced_jurisdiction
        classification["jurisdiction_confidence"] = 1.0

    dtype = classification.get("document_type", "Unknown")
    subtype = classification.get("subtype", "")
    jurisdiction = classification.get("jurisdiction", "Unknown")
    confidence = classification.get("jurisdiction_confidence", 0)
    print(f"Document type : {dtype}{(' — ' + subtype) if subtype else ''}")
    print(f"Jurisdiction  : {jurisdiction}")

    if not jurisdiction or str(jurisdiction).lower() in ("unknown", "unspecified", "none", ""):
        print("\nJurisdiction could not be established from the document.")
        print("Use: check --jurisdiction <country/state/territory>")
        return
    try:
        confidence_value = float(confidence)
    except Exception:
        confidence_value = 0.0
    if confidence_value < 0.60 and not forced_jurisdiction:
        print("Jurisdiction confidence is low; no legal-rule check was performed.")
        print("Use: check --jurisdiction <country/state/territory>")
        return

    queries = classification.get("search_queries") or []
    domains = classification.get("official_domains") or []
    if not queries:
        queries = [
            f"{dtype} {jurisdiction} official legislation regulations requirements",
            f"{dtype} {jurisdiction} official government guidance",
            f"{dtype} {jurisdiction} official court legislation",
        ]

    print("\nSearching authoritative legal sources...")
    urls = []

    # Search engines are useful for discovery, but do not make the check depend
    # on their HTML structure. Add known official starting points first.
    for u in _known_official_legal_urls(dtype, jurisdiction):
        if _legal_host_is_official(u, domains) and u not in urls:
            urls.append(u)

    for q in queries[:5]:
        q2 = f"{q} {jurisdiction} official government legislation"
        for u in _official_web_search(q2, domains, limit=6):
            if u not in urls:
                urls.append(u)
        if len(urls) >= (4 if quick else 8):
            break

    sources = []
    for u in urls[: (5 if quick else 10)]:
        src = _fetch_official_legal_source(u)
        if src:
            sources.append(src)

    if not sources:
        print("\nOfficial source discovery/fetching failed.")
        print("The application could not retrieve authoritative web content from this machine.")
        print("Check your internet/DNS connection and try again.")
        print("No legal conclusions were produced.")
        return

    print(f"✓ Fetched {len(sources)} official source(s).")
    print("✓ Cross-checking document against retrieved rules and guidance...\n")
    report = _legal_build_review(context, classification, sources, quick=quick)

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    base = os.path.splitext(os.path.basename(current_doc_path or document_name()))[0]
    report_path = os.path.join(LEGAL_CHECK_DIR, f"{base}_legal_check_{timestamp}.txt")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write("AI DOCUMENT EDITOR — LEGAL DOCUMENT CHECK\n")
        f.write(f"Document: {document_name()}\n")
        f.write(f"Document type: {dtype}\n")
        f.write(f"Jurisdiction: {jurisdiction}\n")
        f.write(f"Checked: {datetime.datetime.now().isoformat(timespec='seconds')}\n\n")
        f.write(report or "No report returned.")
        f.write("\n\nAUTHORITATIVE SOURCES FETCHED:\n")
        for i, src in enumerate(sources, 1):
            f.write(f"[{i}] {src.get('title','')}\n{src['url']}\n\n")
        f.write("DISCLAIMER: This is an AI-assisted document review, not legal advice or a determination of legal validity.\n")

    print(f"\n[LEGAL CHECK REPORT] {report_path}")


def print_current_document():
    if current_doc_path:
        print(f"[{document_name()}] currently loaded.")
    else:
        print("[No document] No document loaded.")


# =========================================================
# Command parsing
# =========================================================

DOC_COMMANDS_HELP = """
Document commands:
  list docs                                      list .docx files
  open <file.docx>                               load a document
  load <file.docx>                               load/switch document
  current                                        show active document

  inspect                                        analyze document structure (headings, tables, images, headers/footers)
  show                                           show paragraphs + identify headings
  show --all                                     show full document, paragraph indexes + tables
  show tables                                    display table grid layout & cell content

  edit table <T> row <R> col <C> : <text>        edit a specific table cell
  table edit : <natural language instruction>    AI edit tables, e.g. turn all "Low" to "High"
  images                                         index/name embedded images with AI
  replace image <N|imgN> with <file|imgN>        swap image preserving exact drawing formatting

  doc edit : <multiple instructions>             apply all requested formatting/edit operations
                                                 e.g. doc edit : change all headers to Arial green and highlight every Name light blue

  reconstruct <doc.docx/image.png>               rebuild document/image while matching visual formatting
  reconstruct --verify <doc.docx/image.png>      rebuild document and run verification quality check report

  search "notice period"                         search whole sentences
  ask: Who is the landlord?                      ask about loaded document
  <plain question>                                with a document loaded, defaults to document Q&A
  -- <question>                                   ask general AI knowledge outside document context
  summarize                                      summarize the whole document

  change "OLD" to "NEW"                          exact global replacement
  replace "OLD" with "NEW"                      exact global replacement

  rewrite paragraph N in simple language
  rewrite paragraph N: <instruction>             custom AI rewrite
  rewrite paragraph 2 - 3: <instruction>       rewrite a paragraph range

  compare <doc1.docx> <doc2.docx>               compare documents
  improvements                                   understand document + numbered contextual suggestions
  improvements --apply                          apply all applicable suggestions
  check [--quick] [--jurisdiction <place>]   verify legal document against official web sources
  improve format                                improve overall DOCX formatting
  format --Template_name [instructions]          apply ./templates/Template_name.docx formatting
  list templates                                 list templates
  autosave                                       turn automatic change snapshots ON
  autosave off                                   turn automatic change snapshots OFF
  undo                                           undo the last document change
  redo                                           redo the last undone change

  reset                                          restore original loaded document
  save                                           save current document
  save as <file.docx>                            save under a new name

  help                                           show this list
  exit                                           quit
"""




# ---------------------------------------------------------------------------
# PDF <-> DOCX CONVERSION
# ---------------------------------------------------------------------------

def _find_conversion_tool(*names):
    for name in names:
        path = shutil.which(name)
        if path:
            return path
    return None


def _convert_docx_to_pdf(input_path, output_path=None):
    """Convert DOCX to PDF using LibreOffice headless."""
    soffice = _find_conversion_tool("libreoffice", "soffice")
    if not soffice:
        print("[CONVERT] LibreOffice was not found.")
        print("Install it first, e.g.: sudo pacman -S libreoffice")
        return False

    input_path = os.path.abspath(input_path)
    if not os.path.isfile(input_path):
        print(f"[CONVERT] File not found: {input_path}")
        return False

    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + ".pdf"
    output_path = os.path.abspath(output_path)
    output_dir = os.path.dirname(output_path) or os.getcwd()
    os.makedirs(output_dir, exist_ok=True)

    # LibreOffice writes the PDF using the input basename.
    try:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf",
             "--outdir", output_dir, input_path],
            capture_output=True,
            text=True,
            timeout=120,
        )
    except subprocess.TimeoutExpired:
        print("[CONVERT] DOCX → PDF timed out.")
        return False
    except Exception as e:
        print(f"[CONVERT] DOCX → PDF failed: {e}")
        return False

    generated = os.path.join(
        output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf"
    )

    if result.returncode != 0 or not os.path.exists(generated):
        details = (result.stderr or result.stdout or "").strip()
        print("[CONVERT] DOCX → PDF failed.")
        if details:
            print(details)
        return False

    if os.path.abspath(generated) != output_path:
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
            os.replace(generated, output_path)
        except Exception as e:
            print(f"[CONVERT] Could not place output PDF: {e}")
            return False

    print(f"[CONVERT] ✓ DOCX → PDF")
    print(f"Output: {output_path}")
    return True


def _convert_pdf_to_docx(input_path, output_path=None):
    """
    Convert PDF to DOCX using LibreOffice.

    This is primarily intended for text-based PDFs. Scanned/image-only PDFs
    may require OCR and will not necessarily reconstruct the original layout.
    """
    soffice = _find_conversion_tool("libreoffice", "soffice")
    if not soffice:
        print("[CONVERT] LibreOffice was not found.")
        print("Install it first, e.g.: sudo pacman -S libreoffice")
        return False

    input_path = os.path.abspath(input_path)
    if not os.path.isfile(input_path):
        print(f"[CONVERT] File not found: {input_path}")
        return False

    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + ".docx"
    output_path = os.path.abspath(output_path)
    output_dir = os.path.dirname(output_path) or os.getcwd()
    os.makedirs(output_dir, exist_ok=True)

    try:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx",
             "--outdir", output_dir, input_path],
            capture_output=True,
            text=True,
            timeout=120,
        )
    except subprocess.TimeoutExpired:
        print("[CONVERT] PDF → DOCX timed out.")
        return False
    except Exception as e:
        print(f"[CONVERT] PDF → DOCX failed: {e}")
        return False

    generated = os.path.join(
        output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".docx"
    )

    if result.returncode != 0 or not os.path.exists(generated):
        details = (result.stderr or result.stdout or "").strip()
        print("[CONVERT] PDF → DOCX failed.")
        if details:
            print(details)
        return False

    if os.path.abspath(generated) != output_path:
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
            os.replace(generated, output_path)
        except Exception as e:
            print(f"[CONVERT] Could not place output DOCX: {e}")
            return False

    print(f"[CONVERT] ✓ PDF → DOCX")
    print(f"Output: {output_path}")
    return True


def handle_convert_command(user_input):
    """
    Supported:
      convert pdf2docx file.pdf
      convert docx2pdf file.docx
      pdf2docx file.pdf
      docx2pdf file.docx
      pdf to docx file.pdf
      docx to pdf file.docx

    Optional output:
      convert pdf2docx input.pdf output.docx
      convert docx2pdf input.docx output.pdf
    """
    raw = user_input.strip()
    low = raw.lower()

    if not (
        low.startswith("convert ")
        or low.startswith("pdf2docx")
        or low.startswith("docx2pdf")
        or low.startswith("pdf to docx")
        or low.startswith("docx to pdf")
    ):
        return False

    parts = raw.split()
    if len(parts) < 2:
        print("Usage:")
        print("  convert pdf2docx <file.pdf> [output.docx]")
        print("  convert docx2pdf <file.docx> [output.pdf]")
        return True

    if low.startswith("convert "):
        if len(parts) < 3:
            print("Usage:")
            print("  convert pdf2docx <file.pdf> [output.docx]")
            print("  convert docx2pdf <file.docx> [output.pdf]")
            return True
        direction = parts[1].lower()
        input_file = parts[2]
        output_file = parts[3] if len(parts) >= 4 else None
    elif low.startswith("pdf2docx"):
        direction = "pdf2docx"
        input_file = parts[1]
        output_file = parts[2] if len(parts) >= 3 else None
    elif low.startswith("docx2pdf"):
        direction = "docx2pdf"
        input_file = parts[1]
        output_file = parts[2] if len(parts) >= 3 else None
    elif low.startswith("pdf to docx"):
        direction = "pdf2docx"
        input_file = parts[3] if len(parts) >= 4 else ""
        output_file = parts[4] if len(parts) >= 5 else None
    else:
        direction = "docx2pdf"
        input_file = parts[3] if len(parts) >= 4 else ""
        output_file = parts[4] if len(parts) >= 5 else None

    if not input_file:
        print("Please provide an input file.")
        return True

    # Allow paths containing spaces when quoted.
    if '"' in raw or "'" in raw:
        import shlex
        try:
            qparts = shlex.split(raw)
            if low.startswith("convert "):
                direction = qparts[1].lower()
                input_file = qparts[2]
                output_file = qparts[3] if len(qparts) >= 4 else None
            elif low.startswith("pdf2docx"):
                input_file = qparts[1]
                output_file = qparts[2] if len(qparts) >= 3 else None
            elif low.startswith("docx2pdf"):
                input_file = qparts[1]
                output_file = qparts[2] if len(qparts) >= 3 else None
            elif low.startswith("pdf to docx"):
                input_file = qparts[3]
                output_file = qparts[4] if len(qparts) >= 5 else None
            elif low.startswith("docx to pdf"):
                input_file = qparts[3]
                output_file = qparts[4] if len(qparts) >= 5 else None
        except ValueError as e:
            print(f"Invalid quoted path: {e}")
            return True

    direction = direction.replace("-", "").replace("_", "")
    if direction in ("pdf2docx", "pdftodocx"):
        _convert_pdf_to_docx(input_file, output_file)
    elif direction in ("docx2pdf", "docxtopdf"):
        _convert_docx_to_pdf(input_file, output_file)
    else:
        print("Unknown conversion.")
        print("Use: pdf2docx or docx2pdf")
    return True


def handle_doc_command(text):
    t = text.strip()
    low = t.lower()

    if low in ("help", "commands"):
        print(DOC_COMMANDS_HELP)
        return True

    if low in ("list docs", "list documents"):
        files = list_docx_files()
        print("\n".join(files) if files else "No .docx files found in this folder.")
        return True

    m = re.match(r'^(?:open|load)\s+(.+)$', t, re.IGNORECASE)
    if m:
        load_docx(m.group(1))
        return True

    if low in ("current", "current document", "status"):
        print_current_document()
        return True

    if low == "inspect":
        inspect_document()
        return True

    if low in ("show tables", "list tables"):
        show_tables_detail()
        return True

    if low.startswith("edit table "):
        edit_table_cell(t)
        return True

    if low.startswith("table edit") or low.startswith("edit tables"):
        parts = t.split(":", 1)
        if len(parts) == 2:
            ai_table_edit(parts[1].strip())
        else:
            print('Usage: table edit : turn all "Low" to "High"')
        return True

    if low in ("images", "index images", "image index"):
        index_images_command()
        return True

    if low.startswith("replace image "):
        replace_image_in_doc(t)
        return True

    if low.startswith("doc edit") and ":" in t:
        parts = t.split(":", 1)
        doc_edit(parts[1])
        return True

    if re.match(r'^format\s+--', t, re.I):
        format_from_template(t)
        return True

    if low in ("list templates", "templates"):
        list_templates()
        return True

    if low in ("undo", "u"):
        undo_document()
        return True

    if low in ("redo", "r"):
        redo_document()
        return True

    if low == "autosave":
        globals()["AUTOSAVE_ENABLED"] = True
        print("Autosave is ON.")
        return True

    if low == "autosave off":
        globals()["AUTOSAVE_ENABLED"] = False
        print("Autosave is OFF.")
        return True

    m = re.match(r'^reconstruct\s+(.+)$', t, re.IGNORECASE)
    if m:
        reconstruct_document(m.group(1))
        return True

    if low in ("show", "show paragraphs", "list paragraphs"):
        show_paragraphs()
        return True

    m = re.match(r'^compare\s+(.+?)\s+(.+?)$', t, re.IGNORECASE)
    if m:
        compare_documents(m.group(1), m.group(2))
        return True

    m = re.match(r'^search\s+["\'](.+?)["\']$', t, re.IGNORECASE)
    if not m:
        m = re.match(r'^search\s+(.+)$', t, re.IGNORECASE)
    if m:
        search_document(m.group(1))
        return True

    m = re.match(r'^ask\s*:\s*(.+)$', t, re.IGNORECASE)
    if m:
        ask_document(m.group(1))
        return True

    if low == "check" or low.startswith("check --"):
        check_legal_document(t[len("check"):].strip())
        return True

    if low in ("improvements", "improve", "review improvements"):
        improvements()
        return True

    if low in ("summarize", "summary", "summarise"):
        summarize_document()
        return True

    if low in ("improve format", "improve formatting", "format"):
        improve_format()
        return True

    m = re.match(r'^improvements\s+--apply(?:\s+(.+))?$', t, re.IGNORECASE)
    if m:
        apply_improvements(m.group(1).strip() if m.group(1) else None)
        return True

    if low == "reset":
        reset_document()
        return True

    if low in ("show --all", "show all"):
        show_all_paragraphs()
        return True

    m = re.match(r'^(?:change|replace)\s+"(.+?)"\s+(?:to|with)\s+"(.+?)"$',
                 t, re.IGNORECASE)
    if m:
        find_and_replace(m.group(1), m.group(2))
        return True

    if re.match(r'^(?:replace|change)\s+.+\s+(?:everywhere|throughout)',
                t, re.IGNORECASE):
        smart_global_edit(t)
        return True

    if re.match(r'^change\s+.+\s+to\s+.+$', t, re.IGNORECASE):
        smart_global_edit(t)
        return True

    m = re.match(
        r'^rewrite paragraphs?\s+(\d+)\s*[-–]\s*(\d+)\s*(?::|in)\s*(.+)$',
        t, re.IGNORECASE
    )
    if m:
        rewrite_paragraph_range(
            int(m.group(1)), int(m.group(2)), m.group(3).strip()
        )
        return True

    m = re.match(r'^simplify paragraphs?\s+(\d+)\s*[-–]\s*(\d+)$',
                  t, re.IGNORECASE)
    if m:
        rewrite_paragraph_range(
            int(m.group(1)), int(m.group(2)),
            "Rewrite this in simple, easy-to-understand language."
        )
        return True

    m = re.match(r'^simplify paragraph\s+(\d+)$', t, re.IGNORECASE)
    if m:
        rewrite_paragraph(int(m.group(1)),
                          "Rewrite this in simple, easy-to-understand language.")
        return True

    m = re.match(r'^rewrite paragraph\s+(\d+)\s*:\s*(.+)$', t, re.IGNORECASE)
    if m:
        rewrite_paragraph(int(m.group(1)), m.group(2).strip())
        return True

    m = re.match(r'^rewrite paragraph\s+(\d+)\s+in\s+(.+)$', t, re.IGNORECASE)
    if m:
        rewrite_paragraph(int(m.group(1)), f"Rewrite this in {m.group(2).strip()}.")
        return True

    m = re.match(r'^save(?:\s+as\s+(.+))?$', t, re.IGNORECASE)
    if m:
        save_docx(m.group(1) if m.group(1) else None)
        return True

    return False



COMMAND_NAMES = [
    "help", "list docs", "open ", "load ", "current", "inspect", "show",
    "show --all", "show tables", "edit table ", "table edit : ", "images",
    "replace image ", "doc edit : ", "reconstruct ", "search ", "ask: ",
    "summarize", "change ", "replace ", "rewrite paragraph ", "rewrite paragraphs ",
    "compare ", "check", "check --quick", "check --jurisdiction ", "improvements", "improvements --apply", "improve format",
    "format --", "list templates", "autosave", "autosave off", "undo", "redo", "reset",
    "pdf2docx ", "docx2pdf ", "convert pdf2docx ", "convert docx2pdf ",
    "save", "save as ", "exit"
]


def _autocomplete_options():
    opts = list(COMMAND_NAMES)
    try:
        opts += list_docx_files()
    except Exception:
        pass
    try:
        ensure_working_folders()
        opts += [
            "format --" + os.path.splitext(f)[0]
            for f in os.listdir(TEMPLATES_DIR) if f.lower().endswith(".docx")
        ]
        opts += [
            "replace image " + f
            for f in os.listdir(IMAGES_DIR)
            if f.lower().endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp"))
        ]
    except Exception:
        pass
    return sorted(set(opts), key=str.lower)


def setup_terminal_autocomplete():
    try:
        try:
            import readline
        except ImportError:
            import pyreadline3 as readline

        def completer(text, state):
            buffer = readline.get_line_buffer()
            candidates = [x for x in _autocomplete_options()
                          if x.lower().startswith(buffer.lower())]
            if not candidates and text:
                candidates = [x for x in _autocomplete_options()
                              if x.lower().startswith(text.lower())]
            return candidates[state] if state < len(candidates) else None

        readline.set_completer(completer)
        readline.parse_and_bind("tab: complete")
        return True
    except Exception:
        return False


def ask_default_document(question):
    if current_doc is not None:
        ask_document(question)
    else:
        chat_with_openrouter(question, max_tokens=500)


# =========================================================
# Telegram bot mode
# =========================================================

TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_SESSIONS_DIR = os.path.join(os.getcwd(), "telegram_sessions")
TELEGRAM_SESSIONS = {}


def _telegram_session_dir(chat_id):
    path = os.path.join(TELEGRAM_SESSIONS_DIR, str(chat_id))
    os.makedirs(path, exist_ok=True)
    return path


class TelegramDocumentSession:
    def __init__(self, chat_id, source_name):
        self.chat_id = chat_id
        self.source_name = source_name
        self.root = _telegram_session_dir(chat_id)
        self.work_path = os.path.join(self.root, os.path.basename(source_name))
        if not self.work_path.lower().endswith(".docx"):
            self.work_path = os.path.splitext(self.work_path)[0] + ".docx"
        self.autosave_dir = os.path.join(self.root, "autosave")
        os.makedirs(self.autosave_dir, exist_ok=True)
        self.loaded = False
        self.undo_stack = []
        self.redo_stack = []


@contextlib.contextmanager
def _telegram_editor_context(session):
    global current_doc, current_doc_path, original_doc, AUTOSAVE_CURRENT_PATH
    global AUTOSAVE_DIR, AUTOSAVE_ENABLED, UNDO_STACK, REDO_STACK
    import copy
    saved = (current_doc, current_doc_path, original_doc,
             AUTOSAVE_CURRENT_PATH, AUTOSAVE_DIR, AUTOSAVE_ENABLED,
             UNDO_STACK, REDO_STACK)
    try:
        current_doc = Document(session.work_path) if session.loaded and os.path.exists(session.work_path) else None
        original_doc = copy.deepcopy(current_doc) if current_doc is not None else None
        current_doc_path = session.work_path
        AUTOSAVE_CURRENT_PATH = None
        AUTOSAVE_DIR = session.autosave_dir
        AUTOSAVE_ENABLED = True
        UNDO_STACK = session.undo_stack
        REDO_STACK = session.redo_stack
        yield
        session.undo_stack = UNDO_STACK
        session.redo_stack = REDO_STACK
    finally:
        (current_doc, current_doc_path, original_doc,
         AUTOSAVE_CURRENT_PATH, AUTOSAVE_DIR, AUTOSAVE_ENABLED,
         UNDO_STACK, REDO_STACK) = saved


class _TelegramCapture:
    def __init__(self):
        self.parts = []
    def write(self, s):
        self.parts.append(str(s))
        return len(s)
    def flush(self):
        pass
    def getvalue(self):
        return "".join(self.parts)


def _telegram_execute(session, user_input):
    if Document is None:
        return "python-docx is not installed.", False
    low = user_input.strip().lower()
    if re.match(r"^(open|load)\s+", low):
        return "Please upload a DOCX/PDF document instead of using a server file path.", False
    with _telegram_editor_context(session):
        before = current_doc._element.xml if current_doc is not None else None
        capture = _TelegramCapture()
        with contextlib.redirect_stdout(capture):
            try:
                if user_input.startswith("--"):
                    q = user_input[2:].strip()
                    if q:
                        chat_with_openrouter(q, max_tokens=500)
                elif handle_convert_command(user_input):
                    pass
                elif handle_doc_command(user_input):
                    pass
                elif current_doc is not None:
                    if _natural_prompt_is_likely_edit(user_input):
                        natural_document_edit(user_input)
                    else:
                        ask_default_document(user_input)
                else:
                    chat_with_openrouter(user_input, max_tokens=300)
            except Exception as e:
                print(f"Error: {e}")
        after = current_doc._element.xml if current_doc is not None else None
        changed = before != after
        if changed and current_doc is not None:
            current_doc.save(session.work_path)
            autosave_current_document("telegram change")
            session.loaded = True
        return capture.getvalue().strip(), changed


async def _telegram_start(update, context):
    await update.message.reply_text(
        "AI Document Editor is ready. Upload a .docx or .pdf, then send your editing instruction.\n\n"
        "Examples:\n"
        "change Priya to Mayank\n"
        "doc edit : make all headings green and Arial\n"
        "table edit : turn all Low values to High\n"
        "ask: Who is the landlord?\n"
        "-- what is ASL?\n"
        "undo / redo\n"
        "check\n\n"
        "After a document change, the modified DOCX is returned automatically."
    )


async def _telegram_help(update, context):
    await _telegram_start(update, context)


async def _telegram_document(update, context):
    message = update.message
    chat_id = message.chat_id
    tg_doc = message.document
    filename = tg_doc.file_name or "document.docx"
    lower = filename.lower()
    if not lower.endswith((".docx", ".pdf")):
        await message.reply_text("Please upload a .docx or .pdf document.")
        return
    session = TelegramDocumentSession(chat_id, filename)
    TELEGRAM_SESSIONS[chat_id] = session
    try:
        tg_file = await tg_doc.get_file()
        downloaded = os.path.join(session.root, filename)
        await tg_file.download_to_drive(downloaded)
        if lower.endswith(".pdf"):
            converted = os.path.join(session.root, os.path.splitext(filename)[0] + ".docx")
            if not _convert_pdf_to_docx(downloaded, converted):
                await message.reply_text("PDF → DOCX conversion failed. Make sure LibreOffice is installed.")
                return
            session.work_path = converted
        Document(session.work_path)  # validate
        session.loaded = True
        with _telegram_editor_context(session):
            load_docx(session.work_path)
        await message.reply_text(
            f"✓ Loaded {filename}.\nSend your instruction and I will return the modified .docx."
        )
    except Exception as e:
        await message.reply_text(f"Upload failed: {e}")


async def _telegram_text(update, context):
    message = update.message
    chat_id = message.chat_id
    text = (message.text or "").strip()
    session = TELEGRAM_SESSIONS.get(chat_id)
    if not session or not session.loaded:
        await message.reply_text("Please upload a DOCX/PDF document first.")
        return
    await context.bot.send_chat_action(chat_id=chat_id, action="typing")
    result, changed = await asyncio.to_thread(_telegram_execute, session, text)
    if result:
        if len(result) > 3900:
            result = result[:3900] + "\n…(output truncated)"
        await message.reply_text(result)
    if changed and os.path.exists(session.work_path):
        with open(session.work_path, "rb") as f:
            await message.reply_document(document=f, filename=os.path.basename(session.work_path),
                                         caption="✓ Modified document")


async def _telegram_error(update, context):
    print(f"[TELEGRAM] {context.error}")
    if update and getattr(update, "effective_message", None):
        try:
            await update.effective_message.reply_text("An error occurred while processing the document.")
        except Exception:
            pass


def run_telegram_bot():
    if not TELEGRAM_BOT_TOKEN:
        print("TELEGRAM_BOT_TOKEN is not set.")
        print("Set it with: export TELEGRAM_BOT_TOKEN='YOUR_BOT_TOKEN'")
        return
    try:
        from telegram.ext import Application, CommandHandler, MessageHandler, filters
    except ImportError:
        print("Install Telegram support with: pip install python-telegram-bot")
        return
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    application.add_handler(CommandHandler("start", _telegram_start))
    application.add_handler(CommandHandler("help", _telegram_help))
    application.add_handler(MessageHandler(filters.Document.ALL, _telegram_document))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, _telegram_text))
    application.add_error_handler(_telegram_error)
    print("=== AI Document Editor — Telegram Bot ===")
    print("Bot is running. Upload a DOCX/PDF in Telegram.")
    application.run_polling(allowed_updates=None)

# =========================================================
# Main loop
# =========================================================

def main():
    ensure_working_folders()
    setup_terminal_autocomplete()
    print("=== AI Document Editor ===")
    print("Autosave: ON")
    print("Type 'help' for commands.")
    print("With a document loaded, plain questions use document context.")
    print("Use '-- ' before a question for general AI knowledge. Type 'exit' to quit.\n")

    while True:
        try:
            user_input = input("You: ").strip()
            if not user_input:
                continue
            if user_input.lower() == "exit":
                print("Exiting...")
                break

            if user_input.startswith("--"):
                question = user_input[2:].strip()
                if question:
                    chat_with_openrouter(question, max_tokens=500)
                    print()
                continue

            # Capture one pre-edit snapshot for every document-mutating command.
            # Read-only commands and settings changes do not enter history.
            command_low = user_input.lower().strip()
            mutating = (
                command_low.startswith("doc edit")
                or command_low.startswith("table edit")
                or command_low.startswith("edit table ")
                or command_low.startswith("replace image ")
                or command_low.startswith("change ")
                or command_low.startswith("replace ")
                or command_low.startswith("rewrite paragraph")
                or command_low.startswith("rewrite paragraphs")
                or command_low.startswith("simplify paragraph")
                or command_low.startswith("simplify paragraphs")
                or command_low in ("reset", "improve format", "improve formatting", "format")
                or command_low.startswith("format --")
                or command_low.startswith("reconstruct ")
                or command_low.startswith("improvements --apply")
            )
            # Unknown natural-language input can also mutate the document, so
            # snapshot whenever a document is loaded. Read-only questions create
            # no history entry because their XML remains unchanged.
            pre_edit_snapshot = _snapshot_document() if current_doc is not None else None

            if handle_convert_command(user_input):
                print()
                continue

            if handle_doc_command(user_input):
                if mutating and pre_edit_snapshot is not None and current_doc is not None:
                    # A command may have failed internally; only add the snapshot when
                    # the resulting document actually differs from the pre-edit state.
                    try:
                        before_xml = pre_edit_snapshot._element.xml
                        after_xml = current_doc._element.xml
                    except Exception:
                        before_xml = after_xml = None
                    if before_xml != after_xml:
                        _push_undo_snapshot(pre_edit_snapshot)
                print()
                continue

            if current_doc is not None:
                if _natural_prompt_is_likely_edit(user_input):
                    natural_document_edit(user_input)
                    if pre_edit_snapshot is not None:
                        try:
                            if pre_edit_snapshot._element.xml != current_doc._element.xml:
                                _push_undo_snapshot(pre_edit_snapshot)
                        except Exception:
                            pass
                else:
                    ask_default_document(user_input)
                print()
                continue

            response = chat_with_openrouter(user_input, max_tokens=300)
            clean_response(response)
            print("\n")
        except KeyboardInterrupt:
            print("\nExiting...")
            break
        except EOFError:
            print("\nExiting...")
            break


if __name__ == "__main__":
    import sys
    if "--telegram" in sys.argv or (len(sys.argv) > 1 and sys.argv[1].lower() == "telegram"):
        run_telegram_bot()
    else:
        main()
